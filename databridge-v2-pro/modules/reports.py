"""Dynamic report and summary helpers for DataBridge Pro.

The templates bundled with the app are placeholders only.  Reports are generated
from the active dataframe and the selected period (single month / range / all).
"""
from __future__ import annotations

import io
import os
import re
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd


# ────────────────────────────────────────────────────────────────
# Normalization helpers
# ────────────────────────────────────────────────────────────────
def norm_ar(value: Any) -> str:
    s = "" if value is None else str(value).strip()
    s = re.sub(r"\s+", " ", s)
    s = s.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا")
    s = s.replace("ة", "ه").replace("ى", "ي")
    s = s.replace("ـ", "")
    return s.lower()


def is_blank(value: Any) -> bool:
    if pd.isna(value):
        return True
    return str(value).strip().lower() in {"", "nan", "none", "null", "-", "--"}


def is_yes(value: Any) -> bool:
    return norm_ar(value) in {"نعم", "ايوه", "ايوا", "اه", "yes", "y", "true", "1"}


def is_male(value: Any) -> bool:
    return norm_ar(value).replace(" ", "") in {"ذكر", "male", "m"}


def is_female(value: Any) -> bool:
    return norm_ar(value).replace(" ", "") in {"انثي", "female", "f"}


def normalize_result(value: Any) -> str:
    n = norm_ar(value)
    if n in {"ايجابي", "positive", "+", "pos"}:
        return "positive"
    if n in {"سلبي", "negative", "-", "neg"}:
        return "negative"
    return "" if is_blank(value) else n


def num(value: Any) -> int:
    if is_blank(value):
        return 0
    try:
        return int(float(str(value).replace(",", "")))
    except Exception:
        return 0




def numeric_sum(frame: pd.DataFrame, col: Optional[str]) -> float:
    if not col or col not in frame.columns:
        return 0.0
    return float(pd.to_numeric(frame[col], errors="coerce").fillna(0).sum())


def find_best_quantity_col(df: pd.DataFrame, keyword: str, followup_n: Optional[int] = None) -> Optional[str]:
    """Find the actual numeric quantity column, not the Yes/No need/provision column.

    Source files often contain several columns with the same commodity name:
    need? / provided? / quantity. The old logic could select a Yes/No column,
    which made quantities export as zero. We now prefer the candidate with the
    strongest numeric signal within the selected base/follow-up group.
    """
    nk = norm_ar(keyword)
    candidates: List[Tuple[float, int, str]] = []
    bad_tokens = ["تاريخ", "نتيجه", "نتيجة", "زهري", "دعم", "ميثادون"]
    yesno_tokens = ["هل", "احتياج", "يحتاج", "بحاج", "تم تقديم", "تقديم", "استلم", "حصل"]
    qty_tokens = ["كميه", "كمية", "عدد", "qty", "quantity", "no."]
    for c in df.columns:
        nc = norm_ar(c)
        if nk not in nc:
            continue
        if followup_n is None:
            if "متابع" in nc:
                continue
        else:
            if "متابع" not in nc or str(followup_n) not in nc:
                continue
        if any(norm_ar(t) in nc for t in bad_tokens):
            continue
        total = numeric_sum(df, c)
        score = 0
        if total > 0:
            score += 100
        if any(norm_ar(t) in nc for t in qty_tokens):
            score += 25
        if any(norm_ar(t) in nc for t in yesno_tokens):
            score -= 40
        # Short canonical names like "سرنجات" from the converted sheet are valid.
        if nc in {"سرنجات", "واقيات", "مزلقات"} or nc == nk:
            score += 20
        candidates.append((total, score, c))
    if not candidates:
        return None
    # Prefer positive numeric totals; if all are zero, return the best-scored candidate
    # so the output remains deterministic.
    candidates.sort(key=lambda x: (x[0] > 0, x[1], x[0]), reverse=True)
    return candidates[0][2]


def find_confirm_result_col(df: pd.DataFrame) -> Optional[str]:
    candidates: List[Tuple[int, str]] = []
    for c in df.columns:
        nc = norm_ar(c)
        if ("تاكيدي" in nc or "تأكيدي" in nc or "confirm" in nc) and "متابع" not in nc:
            score = 0
            if "نتيجه" in nc or "نتيجة" in nc or "result" in nc:
                score += 100
            if "احاله" in nc or "إحالة" in nc or "referral" in nc:
                score -= 80
            candidates.append((score, c))
    if not candidates:
        return None
    candidates.sort(reverse=True, key=lambda x: x[0])
    return candidates[0][1]


def find_treatment_referral_col(df: pd.DataFrame) -> Optional[str]:
    candidates: List[Tuple[int, str]] = []
    for c in df.columns:
        nc = norm_ar(c)
        if not any(t in nc for t in ["احاله", "الاحاله", "referral", "linked"]):
            continue
        if "متابع" in nc:
            continue
        score = 0
        if any(t in nc for t in ["علاج", "arv", "treatment", "care"]):
            score += 100
        if any(t in nc for t in ["طبيه", "طبي", "medical", "ميثادون", "ost", "زهري", "syphilis"]):
            score -= 100
        candidates.append((score, c))
    if not candidates:
        return None
    candidates.sort(reverse=True, key=lambda x: x[0])
    # Avoid using unrelated referral columns such as medical/OST referrals.
    return candidates[0][1] if candidates[0][0] > 0 else None


def find_col(df: pd.DataFrame, include: List[str], exclude: Optional[List[str]] = None) -> Optional[str]:
    exclude = exclude or []
    for c in df.columns:
        nc = norm_ar(c)
        if all(norm_ar(w) in nc for w in include) and not any(norm_ar(w) in nc for w in exclude):
            return c
    return None


def followup_date_cols(df: pd.DataFrame) -> List[Tuple[int, str]]:
    out: List[Tuple[int, str]] = []
    bad = ["واقي", "مزلق", "زهري", "دعم", "سرنج", "ميثادون", "نتيجه", "نتيجة"]
    for n in range(1, 6):
        for c in df.columns:
            nc = norm_ar(c)
            if norm_ar("زياره متابعه") in nc and str(n) in nc and not any(norm_ar(x) in nc for x in bad):
                out.append((n, c))
                break
    return out


def find_followup_col(df: pd.DataFrame, keyword: str, n: int) -> Optional[str]:
    nk = norm_ar(keyword)
    for c in df.columns:
        nc = norm_ar(c)
        if nk in nc and "متابع" in nc and str(n) in nc:
            return c
    return None


def _parsed_months(series: pd.Series) -> pd.Series:
    """Parse month values consistently for filters.

    Some Streamlit/Pandas paths can coerce one valid-looking source date to
    NaT depending on the cell type.  Keep this helper centralized so the UI,
    summaries and donor reports follow the same date logic.
    """
    dt = pd.to_datetime(series, errors="coerce")
    # A second day-first pass recovers common text dates like 1/6/2026 without
    # changing already parsed ISO/Excel datetime values.
    if dt.isna().any():
        dt2 = pd.to_datetime(series, errors="coerce", dayfirst=True)
        dt = dt.fillna(dt2)
    return dt.dt.to_period("M").astype(str)


def period_mask(series: pd.Series, kind: str, selected_month: Optional[str] = None,
                from_month: Optional[str] = None, to_month: Optional[str] = None) -> pd.Series:
    """Period mask for dated follow-up columns.

    For follow-up columns, blank dates must not be counted. Therefore the
    fallback for all-data mode is dt.notna(), not all rows.
    """
    dt = pd.to_datetime(series, errors="coerce")
    if dt.isna().any():
        dt2 = pd.to_datetime(series, errors="coerce", dayfirst=True)
        dt = dt.fillna(dt2)
    if kind == "single" and selected_month:
        return dt.dt.to_period("M").astype(str).eq(selected_month)
    if kind == "range" and from_month and to_month:
        per = dt.dt.to_period("M").astype(str)
        return per.ge(from_month) & per.le(to_month)
    return dt.notna()


def base_visit_period_mask(series: pd.Series, kind: str, selected_month: Optional[str] = None,
                           from_month: Optional[str] = None, to_month: Optional[str] = None) -> pd.Series:
    """Period mask for basic/reach visits.

    In all-data mode, every loaded base row is counted as reach even if the
    visit-date cell is temporarily unparsed. For a full available range
    (e.g. Jan→Jun on a Jan–Jun upload), rows with unparsed dates are also kept
    so the total matches the uploaded row count. This fixes the 904 vs 905
    issue without counting blank follow-up dates.
    """
    dt = pd.to_datetime(series, errors="coerce")
    if dt.isna().any():
        dt2 = pd.to_datetime(series, errors="coerce", dayfirst=True)
        dt = dt.fillna(dt2)
    if kind == "single" and selected_month:
        return dt.dt.to_period("M").astype(str).eq(selected_month)
    if kind == "range" and from_month and to_month:
        per = dt.dt.to_period("M").astype(str)
        mask = per.ge(from_month) & per.le(to_month)
        valid_months = sorted(per[dt.notna()].unique().tolist())
        # If the user selected the full available period, keep any unparsed
        # base-date row rather than silently dropping one beneficiary.
        if valid_months and from_month <= valid_months[0] and to_month >= valid_months[-1]:
            mask = mask | dt.isna()
        return mask
    return pd.Series(True, index=series.index)


def period_label(kind: str, selected_month: Optional[str], from_month: Optional[str], to_month: Optional[str]) -> str:
    if kind == "single" and selected_month:
        return selected_month
    if kind == "range" and from_month and to_month:
        return f"{from_month} إلى {to_month}"
    return "كل الفترة"


def english_period_label(kind: str, selected_month: Optional[str], from_month: Optional[str], to_month: Optional[str]) -> str:
    if kind == "single" and selected_month:
        return selected_month
    if kind == "range" and from_month and to_month:
        return f"{from_month} to {to_month}"
    return "All Available Period"


# ────────────────────────────────────────────────────────────────
# Area normalization
# ────────────────────────────────────────────────────────────────
def unify_area_name(value: Any) -> str:
    s = norm_ar(value)
    if not s:
        return "غير محدد"
    if any(k in s for k in ["اكتوبر", "دهشور", "كفراوي", "الحي الثاني", "محور الكفراوي"]):
        return "أكتوبر"
    if "وراق" in s:
        return "الوراق السواحل"
    if any(k in s for k in ["دقي", "الدقي", "عمر بن الخطاب", "عمر ابن الخطاب", "مقار"]):
        return "الدقي"
    if "امباب" in s or "البوهي" in s:
        return "إمبابة"
    if "حراني" in s:
        return "الحرانية"
    if "بين السرايات" in s:
        return "بين السرايات"
    if "هرم" in s:
        return "الهرم"
    if "منشاه" in s or "منشا" in s or "القناطر" in s:
        return "طريق منشأة القناطر"
    return str(value).strip()


# ────────────────────────────────────────────────────────────────
# Metric calculation
# ────────────────────────────────────────────────────────────────
def build_report_context(
    full_df: pd.DataFrame,
    month_filter_kind: str = "all",
    selected_month: Optional[str] = None,
    from_month: Optional[str] = None,
    to_month: Optional[str] = None,
) -> Dict[str, Any]:
    df = full_df.copy()
    main_date = find_col(df, ["تاريخ الزياره"], ["متابع"])
    if not main_date:
        main_date = find_col(df, ["تاريخ الزيارة"], ["متابع"])
    if not main_date:
        raise ValueError("لم يتم العثور على عمود تاريخ الزيارة")

    base_mask = base_visit_period_mask(df[main_date], month_filter_kind, selected_month, from_month, to_month)
    base_df = df.loc[base_mask].copy()

    gender_col = find_col(df, ["النوع"]) or find_col(df, ["الجنس"])
    age_col = find_col(df, ["السن"]) or find_col(df, ["العمر"]) or find_col(df, ["age"])
    area_col = find_col(df, ["منطقه"]) or find_col(df, ["منطقة"])
    res_col = find_col(df, ["نتيجه التحليل"], ["متابع", "تاكيدي", "تأكيدي"]) or find_col(df, ["نتيجة التحليل"], ["متابع", "تاكيدي", "تأكيدي"])
    confirm_col = find_confirm_result_col(df)
    referral_col = find_treatment_referral_col(df)
    syr_col = find_best_quantity_col(df, "سرنجات")
    cond_col = find_best_quantity_col(df, "واقيات")
    lube_col = find_best_quantity_col(df, "مزلقات")
    syph_col = find_col(df, ["زهري"], ["متابع"])
    psycho_col = find_col(df, ["دعم نفسي"], ["متابع"])
    meth_col = find_col(df, ["ميثادون"], ["متابع"])
    medical_col = (find_col(df, ["احاله", "طبيه"], ["متابع"]) or find_col(df, ["إحالة", "طبية"], ["متابع"]) or find_col(df, ["medical", "referral"], ["follow"]))

    def sum_col(frame: pd.DataFrame, col: Optional[str]) -> int:
        return int(pd.to_numeric(frame[col], errors="coerce").fillna(0).sum()) if col else 0

    def count_yes(frame: pd.DataFrame, col: Optional[str]) -> int:
        return int(frame[col].apply(is_yes).sum()) if col else 0

    def qty_by_gender(frame: pd.DataFrame, col: Optional[str], want_male: bool) -> int:
        if not col or not gender_col:
            return 0
        mask = frame[gender_col].apply(is_male if want_male else is_female)
        return sum_col(frame.loc[mask], col)

    def count_yes_by_gender(frame: pd.DataFrame, col: Optional[str], want_male: bool) -> int:
        if not col or not gender_col:
            return 0
        mask = frame[gender_col].apply(is_male if want_male else is_female)
        return count_yes(frame.loc[mask], col)

    base_visits = int(len(base_df))
    male = int(base_df[gender_col].apply(is_male).sum()) if gender_col else 0
    female = int(base_df[gender_col].apply(is_female).sum()) if gender_col else 0

    base_results = base_df[res_col].apply(normalize_result) if res_col else pd.Series([], dtype=str)
    basic_tests = int(base_results.isin(["positive", "negative"]).sum())
    basic_positive = int(base_results.eq("positive").sum())
    basic_negative = int(base_results.eq("negative").sum())
    no_result = max(base_visits - basic_tests, 0)
    if gender_col and res_col:
        _gm_base = base_df[gender_col].apply(is_male)
        _gf_base = base_df[gender_col].apply(is_female)
        basic_tests_male = int(base_results[_gm_base].isin(["positive", "negative"]).sum())
        basic_tests_female = int(base_results[_gf_base].isin(["positive", "negative"]).sum())
        basic_positive_male = int(base_results[_gm_base].eq("positive").sum())
        basic_positive_female = int(base_results[_gf_base].eq("positive").sum())
    else:
        basic_tests_male = basic_tests_female = basic_positive_male = basic_positive_female = 0

    confirm_results = base_df[confirm_col].apply(normalize_result) if confirm_col else pd.Series([], dtype=str)
    confirm_positive = int(confirm_results.eq("positive").sum())
    if gender_col and confirm_col:
        confirm_positive_male = int(confirm_results[_gm_base].eq("positive").sum())
        confirm_positive_female = int(confirm_results[_gf_base].eq("positive").sum())
    else:
        confirm_positive_male = confirm_positive_female = 0

    if referral_col:
        referrals = int(base_df[referral_col].apply(is_yes).sum())
    else:
        referrals = confirm_positive

    base_syr = sum_col(base_df, syr_col)
    base_cond = sum_col(base_df, cond_col)
    base_lube = sum_col(base_df, lube_col)

    base_syr_m = qty_by_gender(base_df, syr_col, True)
    base_syr_f = qty_by_gender(base_df, syr_col, False)
    base_cond_m = qty_by_gender(base_df, cond_col, True)
    base_cond_f = qty_by_gender(base_df, cond_col, False)
    base_lube_m = qty_by_gender(base_df, lube_col, True)
    base_lube_f = qty_by_gender(base_df, lube_col, False)

    syphilis = count_yes(base_df, syph_col)
    psycho = count_yes(base_df, psycho_col)
    methadone = count_yes(base_df, meth_col)
    medical_referral = count_yes(base_df, medical_col)

    syphilis_m = count_yes_by_gender(base_df, syph_col, True)
    syphilis_f = count_yes_by_gender(base_df, syph_col, False)
    psycho_m = count_yes_by_gender(base_df, psycho_col, True)
    psycho_f = count_yes_by_gender(base_df, psycho_col, False)
    methadone_m = count_yes_by_gender(base_df, meth_col, True)
    methadone_f = count_yes_by_gender(base_df, meth_col, False)
    medical_m = count_yes_by_gender(base_df, medical_col, True)
    medical_f = count_yes_by_gender(base_df, medical_col, False)

    followup_visits = 0
    followup_male = 0
    followup_female = 0
    followup_tests = 0
    followup_tests_male = 0
    followup_tests_female = 0
    followup_positive = 0
    followup_positive_male = 0
    followup_positive_female = 0
    followup_negative = 0
    fu_syr = fu_cond = fu_lube = 0
    fu_syr_m = fu_syr_f = fu_cond_m = fu_cond_f = fu_lube_m = fu_lube_f = 0
    fu_syphilis = fu_psycho = fu_methadone = 0
    fu_syphilis_m = fu_syphilis_f = fu_psycho_m = fu_psycho_f = fu_methadone_m = fu_methadone_f = 0
    geo_rows: List[Dict[str, Any]] = []

    if area_col:
        for _, r in base_df.iterrows():
            geo_rows.append({"area": unify_area_name(r.get(area_col)), "type": "base"})

    for n, fdate_col in followup_date_cols(df):
        fmask = period_mask(df[fdate_col], month_filter_kind, selected_month, from_month, to_month)
        if not fmask.any():
            continue
        fpart = df.loc[fmask].copy()
        followup_visits += int(fmask.sum())
        if gender_col:
            followup_male += int(fpart[gender_col].apply(is_male).sum())
            followup_female += int(fpart[gender_col].apply(is_female).sum())
        if area_col:
            for _, r in fpart.iterrows():
                geo_rows.append({"area": unify_area_name(r.get(area_col)), "type": "followup"})

        fres_col = find_followup_col(df, "نتيجه تحليل", n) or find_followup_col(df, "نتيجة تحليل", n)
        if fres_col:
            fres = fpart[fres_col].apply(normalize_result)
            followup_tests += int(fres.isin(["positive", "negative"]).sum())
            followup_positive += int(fres.eq("positive").sum())
            followup_negative += int(fres.eq("negative").sum())
            if gender_col:
                _gm_fu = fpart[gender_col].apply(is_male)
                _gf_fu = fpart[gender_col].apply(is_female)
                followup_tests_male += int(fres[_gm_fu].isin(["positive", "negative"]).sum())
                followup_tests_female += int(fres[_gf_fu].isin(["positive", "negative"]).sum())
                followup_positive_male += int(fres[_gm_fu].eq("positive").sum())
                followup_positive_female += int(fres[_gf_fu].eq("positive").sum())

        fsyr_col = find_best_quantity_col(df, "سرنجات", n)
        fcond_col = find_best_quantity_col(df, "واقيات", n)
        flube_col = find_best_quantity_col(df, "مزلقات", n)
        fsyph_col = find_followup_col(df, "زهري", n)
        fpsy_col = find_followup_col(df, "دعم نفسي", n)
        fmeth_col = find_followup_col(df, "ميثادون", n)

        fu_syr += sum_col(fpart, fsyr_col)
        fu_cond += sum_col(fpart, fcond_col)
        fu_lube += sum_col(fpart, flube_col)
        fu_syphilis += count_yes(fpart, fsyph_col)
        fu_psycho += count_yes(fpart, fpsy_col)
        fu_methadone += count_yes(fpart, fmeth_col)

        if gender_col:
            fu_syphilis_m += count_yes_by_gender(fpart, fsyph_col, True)
            fu_syphilis_f += count_yes_by_gender(fpart, fsyph_col, False)
            fu_psycho_m += count_yes_by_gender(fpart, fpsy_col, True)
            fu_psycho_f += count_yes_by_gender(fpart, fpsy_col, False)
            fu_methadone_m += count_yes_by_gender(fpart, fmeth_col, True)
            fu_methadone_f += count_yes_by_gender(fpart, fmeth_col, False)
            gm = fpart[gender_col].apply(is_male)
            gf = fpart[gender_col].apply(is_female)
            fu_syr_m += sum_col(fpart.loc[gm], fsyr_col)
            fu_syr_f += sum_col(fpart.loc[gf], fsyr_col)
            fu_cond_m += sum_col(fpart.loc[gm], fcond_col)
            fu_cond_f += sum_col(fpart.loc[gf], fcond_col)
            fu_lube_m += sum_col(fpart.loc[gm], flube_col)
            fu_lube_f += sum_col(fpart.loc[gf], flube_col)

    total_positive = basic_positive + followup_positive
    unconfirmed = max(total_positive - confirm_positive, 0)

    if geo_rows:
        geo_df = pd.DataFrame(geo_rows)
        geo_counts = geo_df["area"].value_counts().reset_index()
        geo_counts.columns = ["المنطقة", "العدد"]
    else:
        geo_counts = pd.DataFrame(columns=["المنطقة", "العدد"])

    # Chart-support summaries for report templates
    if main_date and main_date in base_df.columns:
        _mdt = pd.to_datetime(base_df[main_date], errors="coerce")
        monthly_reach = (_mdt.dt.to_period("M").astype(str).value_counts().sort_index().reset_index())
        monthly_reach.columns = ["month", "count"]
    else:
        monthly_reach = pd.DataFrame(columns=["month", "count"])

    if age_col and age_col in base_df.columns:
        age_counts = base_df[age_col].astype(str).str.strip().value_counts().reset_index()
        age_counts.columns = ["age", "count"]
    else:
        age_counts = pd.DataFrame(columns=["age", "count"])

    ctx = {
        "period_kind": month_filter_kind,
        "period_label": period_label(month_filter_kind, selected_month, from_month, to_month),
        "period_label_en": english_period_label(month_filter_kind, selected_month, from_month, to_month),
        "base_visits": base_visits,
        "male": male,
        "female": female,
        "followup_visits": followup_visits,
        "followup_male": followup_male,
        "followup_female": followup_female,
        "total_contacts": base_visits + followup_visits,
        "basic_tests": basic_tests,
        "basic_tests_male": basic_tests_male,
        "basic_tests_female": basic_tests_female,
        "followup_tests": followup_tests,
        "followup_tests_male": followup_tests_male,
        "followup_tests_female": followup_tests_female,
        "total_tests": basic_tests + followup_tests,
        "basic_negative": basic_negative,
        "followup_negative": followup_negative,
        "total_negative": basic_negative + followup_negative,
        "basic_positive": basic_positive,
        "basic_positive_male": basic_positive_male,
        "basic_positive_female": basic_positive_female,
        "followup_positive": followup_positive,
        "followup_positive_male": followup_positive_male,
        "followup_positive_female": followup_positive_female,
        "rapid_positive_total": total_positive,
        "rapid_positive_total_male": basic_positive_male + followup_positive_male,
        "rapid_positive_total_female": basic_positive_female + followup_positive_female,
        "confirm_positive": confirm_positive,
        "confirm_positive_male": confirm_positive_male,
        "confirm_positive_female": confirm_positive_female,
        "referrals_linked": referrals,
        "rapid_positive_unconfirmed": unconfirmed,
        "no_test_result": no_result,
        "base_syringes": base_syr,
        "base_condoms": base_cond,
        "base_lubricants": base_lube,
        "followup_syringes": fu_syr,
        "followup_condoms": fu_cond,
        "followup_lubricants": fu_lube,
        "total_syringes": base_syr + fu_syr,
        "total_condoms": base_cond + fu_cond,
        "total_lubricants": base_lube + fu_lube,
        "base_syringes_male": base_syr_m,
        "base_syringes_female": base_syr_f,
        "base_condoms_male": base_cond_m,
        "base_condoms_female": base_cond_f,
        "base_lubricants_male": base_lube_m,
        "base_lubricants_female": base_lube_f,
        "followup_syringes_male": fu_syr_m,
        "followup_syringes_female": fu_syr_f,
        "followup_condoms_male": fu_cond_m,
        "followup_condoms_female": fu_cond_f,
        "followup_lubricants_male": fu_lube_m,
        "followup_lubricants_female": fu_lube_f,
        "total_syringes_male": base_syr_m + fu_syr_m,
        "total_syringes_female": base_syr_f + fu_syr_f,
        "total_condoms_male": base_cond_m + fu_cond_m,
        "total_condoms_female": base_cond_f + fu_cond_f,
        "total_lubricants_male": base_lube_m + fu_lube_m,
        "total_lubricants_female": base_lube_f + fu_lube_f,
        "syphilis": syphilis + fu_syphilis,
        "syphilis_male": syphilis_m + fu_syphilis_m,
        "syphilis_female": syphilis_f + fu_syphilis_f,
        "psychosocial": psycho + fu_psycho,
        "psychosocial_male": psycho_m + fu_psycho_m,
        "psychosocial_female": psycho_f + fu_psycho_f,
        "methadone_referrals": methadone + fu_methadone,
        "methadone_referrals_male": methadone_m + fu_methadone_m,
        "methadone_referrals_female": methadone_f + fu_methadone_f,
        "medical_referrals": medical_referral,
        "medical_referrals_male": medical_m,
        "medical_referrals_female": medical_f,
        "geo_counts": geo_counts,
        "monthly_reach": monthly_reach,
        "age_counts": age_counts,
        "base_df": base_df,
    }
    return ctx


# ────────────────────────────────────────────────────────────────
# Data summary table and Excel
# ────────────────────────────────────────────────────────────────
def build_summary_display_df(ctx: Dict[str, Any]) -> pd.DataFrame:
    rows = [
        ["إجمالي الزوار", ctx["total_contacts"], ctx["male"] + ctx.get("followup_male", 0), ctx["female"] + ctx.get("followup_female", 0)],
        ["إجمالي الوصول", ctx["base_visits"], ctx["male"], ctx["female"]],
        ["إجمالي التحاليل", ctx["basic_tests"], ctx.get("basic_tests_male", 0), ctx.get("basic_tests_female", 0)],
        ["أعداد إيجابي تحليل سريع", ctx["rapid_positive_total"], ctx.get("rapid_positive_total_male", 0), ctx.get("rapid_positive_total_female", 0)],
        ["أعداد إيجابي تحليل تأكيدي", ctx["confirm_positive"], ctx.get("confirm_positive_male", 0), ctx.get("confirm_positive_female", 0)],
        ["أعداد متابعة", ctx["followup_visits"], ctx.get("followup_male", 0), ctx.get("followup_female", 0)],
        ["أعداد تحليل متابعة", ctx["followup_tests"], ctx.get("followup_tests_male", 0), ctx.get("followup_tests_female", 0)],
        ["أعداد إيجابي متابعة تحليل سريع", ctx["followup_positive"], ctx.get("followup_positive_male", 0), ctx.get("followup_positive_female", 0)],
        ["إجمالي إيجابي متابعة تحليل تأكيدي", 0, 0, 0],
        ["إجمالي توزيع سرنجات زوار", ctx["total_syringes"], ctx["total_syringes_male"], ctx["total_syringes_female"]],
        ["إجمالي توزيع واقيات زوار", ctx["total_condoms"], ctx["total_condoms_male"], ctx["total_condoms_female"]],
        ["إجمالي توزيع مزلقات زوار", ctx["total_lubricants"], ctx["total_lubricants_male"], ctx["total_lubricants_female"]],
        ["إجمالي توزيع سرنجات أول مرة", ctx["base_syringes"], ctx["base_syringes_male"], ctx["base_syringes_female"]],
        ["إجمالي توزيع واقيات أول مرة", ctx["base_condoms"], ctx["base_condoms_male"], ctx["base_condoms_female"]],
        ["إجمالي توزيع مزلقات أول مرة", ctx["base_lubricants"], ctx["base_lubricants_male"], ctx["base_lubricants_female"]],
        ["إجمالي توزيع سرنجات متابعة", ctx["followup_syringes"], ctx["followup_syringes_male"], ctx["followup_syringes_female"]],
        ["إجمالي توزيع واقيات متابعة", ctx["followup_condoms"], ctx["followup_condoms_male"], ctx["followup_condoms_female"]],
        ["إجمالي توزيع مزلقات متابعة", ctx["followup_lubricants"], ctx["followup_lubricants_male"], ctx["followup_lubricants_female"]],
        ["حالات دعم نفسي", ctx["psychosocial"], ctx.get("psychosocial_male", 0), ctx.get("psychosocial_female", 0)],
        ["حالات إحالة طبية", ctx["medical_referrals"], ctx.get("medical_referrals_male", 0), ctx.get("medical_referrals_female", 0)],
        ["حالات إحالة ميثادون", ctx["methadone_referrals"], ctx.get("methadone_referrals_male", 0), ctx.get("methadone_referrals_female", 0)],
    ]
    return pd.DataFrame(rows, columns=["البيان", "الإجمالي", "رجال", "سيدات"])



def _template_dir() -> str:
    """Return templates folder both in dev and PyInstaller/Tauri frozen mode."""
    import sys
    base = getattr(sys, "_MEIPASS", None)
    candidates = []
    if base:
        candidates.append(os.path.join(base, "templates"))
    candidates.extend([
        os.path.join(os.getcwd(), "templates"),
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates"),
    ])
    for c in candidates:
        if os.path.isdir(c):
            return c
    return candidates[0]


def _template_path(filename: str) -> str:
    path = os.path.join(_template_dir(), filename)
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template not found: {path}")
    return path


def _fmt_int(value: Any) -> str:
    try:
        return f"{int(value):,}"
    except Exception:
        return str(value)


def _pct(n: int, d: int, decimals: int = 1) -> str:
    return f"{(n / d * 100):.{decimals}f}%" if d else "0.0%"


def _as_int(value: Any, default: int = 0) -> int:
    try:
        if value is None or value == "":
            return default
        return int(float(str(value).replace(",", "")))
    except Exception:
        return default


def _get_target(ctx: Dict[str, Any], indicator: str, key: Optional[str] = None) -> int:
    targets = ctx.get("kpi_targets") or ctx.get("targets") or {}
    if not isinstance(targets, dict):
        return 0
    key = key or str(ctx.get("target_key") or "")
    if key == "Annual":
        key = "annual"
    if not key:
        key = _infer_target_key(ctx)
    indicator_targets = targets.get(indicator, {})
    if not isinstance(indicator_targets, dict):
        return 0
    return _as_int(indicator_targets.get(key, 0), 0)


def _infer_target_key(ctx: Dict[str, Any]) -> str:
    """Infer the most relevant target bucket when the UI did not pass one."""
    kind = str(ctx.get("period_kind", ""))
    label = str(ctx.get("period_label", ctx.get("period_label_en", "")))
    if kind == "all":
        return "annual"
    months = re.findall(r"2026-(\d{2})", label)
    if months:
        nums = [int(m) for m in months]
        if all(1 <= m <= 3 for m in nums):
            return "Q1"
        if all(4 <= m <= 6 for m in nums):
            return "Q2"
        if all(7 <= m <= 9 for m in nums):
            return "Q3"
        if all(10 <= m <= 12 for m in nums):
            return "Q4"
    return "annual" if kind == "all" else "Q2"


def _target_label(ctx: Dict[str, Any], key: str) -> str:
    return str(ctx.get("target_label") or ("Annual" if key == "annual" else key))


def _achievement(actual: int, target: int) -> str:
    return _pct(actual, target) if target else "N/A"


def _target_or_dash(target: int) -> str:
    return _fmt_int(target) if target else "-"


def _ctx_months(ctx: Dict[str, Any]) -> List[str]:
    """Return selected months present in the report context, sorted as YYYY-MM."""
    months: List[str] = []
    mdf = ctx.get("monthly_reach")
    if isinstance(mdf, pd.DataFrame) and not mdf.empty and "month" in mdf.columns:
        months = [str(x) for x in mdf["month"].dropna().astype(str).tolist() if re.match(r"^\d{4}-\d{2}$", str(x))]
    if not months:
        label = str(ctx.get("period_label_en") or ctx.get("period_label") or "")
        months = re.findall(r"\d{4}-\d{2}", label)
    return sorted(set(months))


def _quarter_from_months(months: List[str]) -> Optional[str]:
    if len(months) != 3:
        return None
    nums = [int(m[-2:]) for m in months]
    if nums == [1, 2, 3]:
        return "Q1"
    if nums == [4, 5, 6]:
        return "Q2"
    if nums == [7, 8, 9]:
        return "Q3"
    if nums == [10, 11, 12]:
        return "Q4"
    return None


def _period_scope(ctx: Dict[str, Any], target_key: str = "") -> Dict[str, str]:
    """Human-readable period labels for M&E report headings.

    The original template is Q2-specific, but exported M&E reports can now be
    monthly, quarterly, year-to-date, annual, or a custom selected period.
    """
    kind = str(ctx.get("period_kind") or "")
    months = _ctx_months(ctx)
    quarter = _quarter_from_months(months)
    target_key_norm = "annual" if target_key in {"Annual", "annual"} else target_key

    if kind == "single":
        short = "Monthly"
        actual = "Monthly"
        perf = "Monthly"
    elif len(months) == 12:
        short = "Annual"
        actual = "Annual"
        perf = "Annual"
    elif quarter:
        short = quarter
        actual = quarter
        perf = quarter
    elif kind == "all" and len(months) < 12:
        short = "Year-to-Date"
        actual = "YTD"
        perf = "Year-to-Date"
    else:
        short = "Selected Period"
        actual = "Selected Period"
        perf = "Selected Period"

    # If the user intentionally selected annual targets for a partial period,
    # keep target label Annual but keep actual label YTD/Selected Period.
    target = "Annual" if target_key_norm == "annual" else (target_key_norm or short)
    return {"short": short, "actual": actual, "performance": perf, "target": target}


def _monthly_trend_text(ctx: Dict[str, Any]) -> str:
    mdf = ctx.get("monthly_reach")
    if not isinstance(mdf, pd.DataFrame) or mdf.empty:
        return ""
    parts = []
    for _, r in mdf.sort_values("month").iterrows():
        parts.append(f"{str(r['month'])}: {int(r['count']):,}")
    return " | ".join(parts)


def _age_narrative(ctx: Dict[str, Any]) -> str:
    adf = ctx.get("age_counts")
    base = int(ctx.get("base_visits", 0) or 0)
    if not isinstance(adf, pd.DataFrame) or adf.empty or base <= 0:
        return "Age distribution was calculated for the selected period."
    work = adf.copy()
    work["age"] = work["age"].astype(str).str.strip()
    work["count"] = pd.to_numeric(work["count"], errors="coerce").fillna(0).astype(int)

    def _age_start(label: str) -> Optional[int]:
        m = re.search(r"(\d+)", label)
        return int(m.group(1)) if m else None

    band_mask = work["age"].apply(lambda x: (_age_start(x) is not None and 20 <= _age_start(x) <= 44))
    band_total = int(work.loc[band_mask, "count"].sum())
    top = work.sort_values("count", ascending=False).iloc[0]
    pct = _pct(band_total, base, 0).replace(".0%", "%")
    return (
        f"The 20–44 age band accounts for the large majority of beneficiaries "
        f"({band_total:,} of {base:,}, ~{pct}), with the {top['age']} group the single largest cohort "
        f"({int(top['count']):,} beneficiaries). Younger and older age groups should continue to receive targeted outreach."
    )


def _month_label_ar(period: str) -> str:
    """Convert YYYY-MM to Arabic display label used by monthly template."""
    names = {
        "01": "يناير", "02": "فبراير", "03": "مارس", "04": "أبريل",
        "05": "مايو", "06": "يونيو", "07": "يوليو", "08": "أغسطس",
        "09": "سبتمبر", "10": "أكتوبر", "11": "نوفمبر", "12": "ديسمبر",
    }
    if re.match(r"^\d{4}-\d{2}$", str(period)):
        y, m = str(period).split("-")
        return f"شهر {names.get(m, m)} {y}"
    return str(period)


def _period_doc_label(ctx: Dict[str, Any]) -> str:
    if ctx.get("period_kind") == "single":
        return _month_label_ar(str(ctx.get("period_label", "")))
    return str(ctx.get("period_label", ""))


def _safe_set_cell(cell, text: Any) -> None:
    cell.text = _fmt_int(text) if isinstance(text, int) else str(text)


def _replace_in_paragraph(paragraph, new_text: str) -> None:
    """Replace paragraph text while preserving the first run's basic style."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def _iter_table_paragraphs(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _all_paragraphs(doc):
    # Body paragraphs
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        yield from _iter_table_paragraphs(tbl)

    # Headers/footers are separate Word parts; template titles often live there.
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            for p in part.paragraphs:
                yield p
            for tbl in part.tables:
                yield from _iter_table_paragraphs(tbl)


def _replace_text_everywhere(doc, replacements: Dict[str, str]) -> None:
    for p in _all_paragraphs(doc):
        txt = p.text
        if not txt:
            continue
        new = txt
        for old, val in replacements.items():
            if old in new:
                new = new.replace(old, val)
        if new != txt:
            _replace_in_paragraph(p, new)


def _set_table_value_by_label(doc, label_contains: str, value: Any, value_col_offset: int = 1) -> bool:
    needle = norm_ar(label_contains)
    for tbl in doc.tables:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if needle and needle in norm_ar(cell.text):
                    j = min(i + value_col_offset, len(row.cells) - 1)
                    _safe_set_cell(row.cells[j], value)
                    return True
    return False


def _set_table_row_values(doc, first_cell_contains: str, values: List[Any]) -> bool:
    needle = norm_ar(first_cell_contains)
    for tbl in doc.tables:
        for row in tbl.rows:
            if row.cells and needle in norm_ar(row.cells[0].text):
                for i, val in enumerate(values, start=1):
                    if i < len(row.cells):
                        _safe_set_cell(row.cells[i], val)
                return True
    return False


def _replace_docx_media(docx_bytes: bytes, image_map: Dict[str, bytes]) -> bytes:
    if not image_map:
        return docx_bytes
    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in image_map:
                data = image_map[item.filename]
            zout.writestr(item, data)
    return dst.getvalue()


def _clean_docx_extended_properties(docx_bytes: bytes, title: str, subject: str) -> bytes:
    """Clean hidden Office extended metadata that python-docx core_properties does not edit.

    Word templates can retain old values in docProps/app.xml (for example,
    TitlesOfParts still showing an old Q2 report title). This function rewrites
    docProps XML after python-docx saves the file and after image replacement.
    """
    from xml.sax.saxutils import escape

    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    safe_title = escape(str(title or "Befrienders Program — M&E Report"))
    safe_subject = escape(str(subject or "M&E Report"))
    titles_of_parts = (
        f"<TitlesOfParts><vt:vector size=\"1\" baseType=\"lpstr\">"
        f"<vt:lpstr>{safe_title}</vt:lpstr>"
        f"</vt:vector></TitlesOfParts>"
    )

    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("docProps/") and item.filename.endswith(".xml"):
                text = data.decode("utf-8", errors="ignore")
                text = text.replace("Be Frienders", "Befrienders")
                text = text.replace("BeFrienders", "Befrienders")
                text = text.replace("Corrected and validated Q2 2026 M&E report", safe_subject)
                text = text.replace("Q2 2026 Quarterly Report", safe_title)
                text = text.replace("Q2 2026", safe_title)
                text = re.sub(r"<TitlesOfParts>.*?</TitlesOfParts>", titles_of_parts, text, flags=re.DOTALL)
                if "<TitlesOfParts>" not in text and "</Properties>" in text and item.filename.endswith("app.xml"):
                    text = text.replace("</Properties>", titles_of_parts + "</Properties>")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    return dst.getvalue()


def _matplotlib_chart_png(kind: str, ctx: Dict[str, Any], width: int, height: int) -> bytes:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np

    fig, ax = plt.subplots(figsize=(width / 100, height / 100), dpi=100)
    title_period = ctx.get("period_label_en") or ctx.get("period_label") or "Selected Period"

    def finish():
        fig.tight_layout()
        bio = io.BytesIO()
        fig.savefig(bio, format="png", dpi=100, bbox_inches="tight")
        plt.close(fig)
        return bio.getvalue()

    if kind == "monthly":
        mdf = ctx.get("monthly_reach")
        if isinstance(mdf, pd.DataFrame) and not mdf.empty:
            x = mdf["month"].astype(str).tolist()
            y = mdf["count"].astype(int).tolist()
        else:
            x = [str(ctx.get("period_label", "Period"))]
            y = [int(ctx.get("base_visits", 0))]
        ax.bar(x, y, color="#2166A8")
        ax.set_title(f"Beneficiaries Reached — {title_period}", fontsize=20, weight="bold")
        ax.set_ylabel("Beneficiaries")
        for i, v in enumerate(y):
            ax.text(i, v, _fmt_int(v), ha="center", va="bottom", fontsize=12, weight="bold")
        return finish()

    if kind == "gender":
        vals = [int(ctx.get("male", 0)), int(ctx.get("female", 0))]
        labels = ["Male", "Female"]
        if sum(vals) == 0:
            vals = [1, 0]
        wedges, texts, autotexts = ax.pie(vals, labels=labels, autopct=lambda p: f"{p:.1f}%", startangle=90, colors=["#2166A8", "#D95732"], wedgeprops=dict(width=0.42))
        for i, t in enumerate(autotexts):
            if sum(vals):
                t.set_text(f"{vals[i]}\n({vals[i] / max(sum(vals), 1) * 100:.1f}%)")
        ax.set_title(f"Gender Distribution — {title_period}", fontsize=18, weight="bold")
        return finish()

    if kind == "age":
        order = ['15 - 19', '20 - 24', '25 - 29', '30 - 34', '35 - 39', '40 - 44', '45 - 49', '50 او اكثر', '50+']
        adf = ctx.get("age_counts")
        counts = {}
        if isinstance(adf, pd.DataFrame) and not adf.empty:
            for _, r in adf.iterrows():
                counts[str(r['age']).strip()] = int(r['count'])
        labels = [x for x in order if x in counts]
        if not labels:
            labels = list(counts.keys()) or ["N/A"]
        y = [counts.get(x, 0) for x in labels]
        ax.bar(labels, y, color="#1FA37A")
        ax.set_title(f"Age Distribution of Beneficiaries — {title_period}", fontsize=18, weight="bold")
        ax.set_ylabel("Beneficiaries")
        ax.tick_params(axis='x', rotation=20)
        for i, v in enumerate(y):
            ax.text(i, v, _fmt_int(v), ha="center", va="bottom", fontsize=11)
        return finish()

    if kind == "geo":
        gdf = ctx.get("geo_counts")
        if isinstance(gdf, pd.DataFrame) and not gdf.empty:
            gdf = gdf.head(10).iloc[::-1]
            labels = gdf["المنطقة"].astype(str).tolist()
            vals = gdf["العدد"].astype(int).tolist()
        else:
            labels, vals = ["N/A"], [0]
        ax.barh(labels, vals, color="#2166A8")
        ax.set_title("التوزيع الجغرافي - أسماء المناطق الموحدة", fontsize=18, weight="bold")
        ax.set_xlabel("عدد المستفيدين والزيارات")
        for i, v in enumerate(vals):
            ax.text(v, i, f" {v:,}", va="center", fontsize=11, weight="bold")
        return finish()

    if kind == "commodities":
        labels = ["Syringes", "Condoms", "Lubricants"]
        basic = [int(ctx.get("base_syringes", 0)), int(ctx.get("base_condoms", 0)), int(ctx.get("base_lubricants", 0))]
        fu = [int(ctx.get("followup_syringes", 0)), int(ctx.get("followup_condoms", 0)), int(ctx.get("followup_lubricants", 0))]
        x = np.arange(len(labels)); w = 0.35
        ax.bar(x - w/2, basic, w, label="Basic visit", color="#1f77b4")
        ax.bar(x + w/2, fu, w, label="Follow-up visit", color="#ff7f0e")
        ax.set_title(f"Commodities Distributed - {title_period}", fontsize=18, weight="bold")
        ax.set_ylabel("Units"); ax.set_xticks(x); ax.set_xticklabels(labels)
        ax.legend()
        for i, v in enumerate(basic):
            ax.text(i - w/2, v, _fmt_int(v), ha="center", va="bottom", fontsize=10, weight="bold")
        for i, v in enumerate(fu):
            ax.text(i + w/2, v, _fmt_int(v), ha="center", va="bottom", fontsize=10, weight="bold")
        return finish()

    return b""


def build_summary_excel_bytes(ctx: Dict[str, Any]) -> bytes:
    """Fill the user's real Excel summary template while preserving layout/styles."""
    from openpyxl import load_workbook

    is_single = ctx.get("period_kind") == "single"
    template_name = "ملخص شهر 6.xlsx" if is_single else "BeFrienders Q2 2026 Summary.xlsx"
    wb = load_workbook(_template_path(template_name))
    ws = wb.active
    ws.title = _month_label_ar(str(ctx.get("period_label", ""))).replace("شهر ", "") if is_single else "Q2 2026"

    # Period/title cell: template has it in A2 or B2 depending on workbook.
    display_label = _month_label_ar(str(ctx.get("period_label", ""))) if is_single else str(ctx.get("period_label", ""))
    for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 4)):
        for cell in row:
            if isinstance(cell.value, str) and ("شهر" in cell.value or "أبريل" in cell.value or "يونيو" in cell.value):
                cell.value = display_label
                break

    summary = build_summary_display_df(ctx)
    values_by_label = {
        str(r["البيان"]): [r["الإجمالي"], r["رجال"], r["سيدات"]]
        for _, r in summary.iterrows()
    }
    for row in ws.iter_rows():
        for cell in row:
            label = str(cell.value).strip() if cell.value is not None else ""
            if label in values_by_label:
                vals = values_by_label[label]
                start_col = cell.column + 1
                for i, val in enumerate(vals):
                    ws.cell(cell.row, start_col + i).value = 0 if pd.isna(val) else int(val)
                break
    out = io.BytesIO(); wb.save(out); return out.getvalue()



def build_monthly_arabic_report_docx(ctx: Dict[str, Any]) -> bytes:
    """Fill the user's official Arabic monthly DOCX form in-place.

    This function intentionally does not rebuild the report. It opens
    templates/التقرير الشهري.docx, updates only existing table cells / value
    paragraphs, and leaves the original form structure, tables, map section,
    spacing and static text in place.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(_template_path("التقرير الشهري.docx"))
    label = _month_label_ar(str(ctx.get("period_label", "")))

    def set_cell(cell, value: Any) -> None:
        _safe_set_cell(cell, _fmt_int(value) if isinstance(value, int) else value)

    # ── Section A: General information table ──
    if len(doc.tables) > 0:
        general = doc.tables[0]
        for row in general.rows:
            tx = row.cells[0].text
            if "الفترة المشمولة" in tx:
                set_cell(row.cells[0], f"الفترة المشمولة بالتقريرالشهري:  {label}")

    # ── Section B: Quantitative data table ──
    # The official monthly form has one large table. Keep it intact and update
    # the Result column only. Commodity rows represent total service volume
    # during the month, so they include both basic and follow-up distributions.
    monthly_values_by_label = {
        "عدد الأشخاص الذين تم الوصول إليهم": ctx["base_visits"],
        "عدد الأشخاص الذين أجروا اختبار": ctx["basic_tests"],
        "ايجابي HIV باستخدام الاختبار السريع": ctx["basic_positive"],
        "ايجابي HIV تأكيدي": ctx["confirm_positive"],
        "تم ربطهم بمراكز العلاج": ctx["referrals_linked"],
        "إجمالي عدد السرنجات التي تم توزيعها": ctx["total_syringes"],
        "إجمالي عدد الواقيات الذكرية التي تم توزيعها": ctx["total_condoms"],
        "إجمالي عدد المزلقات التي تم توزيعها": ctx["total_lubricants"],
        "عدد المستفيدين الذين حصلوا علي معلومات": ctx["total_contacts"],
        "خدمات المشورة والدعم النفسي": ctx["psychosocial"],
        "اختبارات الأمراض المنقولة": ctx["syphilis"],
        "خدمات الحد من الضرر": ctx["methadone_referrals"],
        "عدد المتابعات المقامة": ctx["followup_visits"],
    }
    if len(doc.tables) > 1:
        qtbl = doc.tables[1]
        for row in qtbl.rows:
            if len(row.cells) < 4:
                continue
            indicator = row.cells[2].text
            for needle, value in monthly_values_by_label.items():
                if norm_ar(needle) in norm_ar(indicator):
                    set_cell(row.cells[3], value)
                    break

    # ── Narrative: update only existing paragraphs that already contain values.
    # Do not remove sections such as challenges, success stories, objectives,
    # notes, or the field map table.
    month_name = label.replace("شهر ", "")
    avg_syr = round(ctx["total_syringes"] / max(ctx["base_visits"], 1))
    avg_cond = round(ctx["total_condoms"] / max(ctx["base_visits"], 1))
    avg_lube = round(ctx["total_lubricants"] / max(ctx["base_visits"], 1))

    for p in doc.paragraphs:
        tx = p.text.strip()
        if not tx:
            continue
        ntx = norm_ar(tx)
        new_text = None
        if "التقرير الشهري" in ntx and "2026" in tx:
            new_text = f"التقرير الشهري {month_name}"
        elif "اجمالي عدد المستفيدين الذين تم الوصول اليهم" in ntx:
            new_text = f"إجمالي عدد المستفيدين الذين تم الوصول إليهم: {ctx['base_visits']:,} مستفيدا ."
        elif "عدد الرجال" in ntx and "عدد السيدات" in ntx:
            new_text = f"عدد الرجال: {ctx['male']:,} مستفيدا ."
            # Keep the original paragraph one-line style by adding female count.
            new_text += f"\nعدد السيدات: {ctx['female']:,} مستفيدة ."
        elif "عدد زيارات المتابعه" in ntx or "عدد زيارات المتابعة" in tx:
            new_text = f"عدد زيارات المتابعة: {ctx['followup_visits']:,} زيارة متابعة ."
        elif "عدد المستفيدين الذين وافقوا علي اجراء الفحص" in ntx or "عدد المستفيدين الذين وافقوا على إجراء الفحص" in tx:
            new_text = f"عدد المستفيدين الذين وافقوا على إجراء الفحص: {ctx['basic_tests']:,} مستفيدا ."
        elif "عدد النتائج الايجابيه المكتشفه" in ntx or "عدد النتائج الإيجابية المكتشفة" in tx:
            new_text = f"عدد النتائج الإيجابية المكتشفة: {ctx['basic_positive']:,} حالة ايجابية تحليل سريع."
        elif "عدد النتائج في التحليل التاكيد" in ntx:
            new_text = f"عدد النتائج في التحليل التاكيدى {ctx['confirm_positive']:,} ."
        elif "عدد اختبارات الزهري المنفذه" in ntx or "عدد اختبارات الزهري المنفذة" in tx:
            new_text = f"عدد اختبارات الزهري المنفذة: {ctx['syphilis']:,} اختبارات ."
        elif "عدد جلسات الدعم النفسي الفردي" in ntx:
            new_text = f"عدد جلسات الدعم النفسي الفردي: {ctx['psychosocial']:,} جلسات ."
        elif "اجمالي السرنجات الموزعه" in ntx:
            new_text = f"إجمالي السرنجات الموزعة: {ctx['total_syringes']:,} سرنجة."
        elif "متوسط التوزيع" in ntx and "سرنج" in ntx:
            new_text = f"متوسط التوزيع: {avg_syr:,} سرنجة لكل مستفيد ."
        elif "اجمالي الواقيات الذكريه الموزعه" in ntx:
            new_text = f"إجمالي الواقيات الذكرية الموزعة: {ctx['total_condoms']:,} واقيا."
        elif "واقي" in ntx and "متوسط التوزيع" in ntx:
            new_text = f"متوسط التوزيع: {avg_cond:,} واقيا لكل مستفيد ."
        elif "اجمالي المزلقات الموزعه" in ntx:
            new_text = f"إجمالي المزلقات الموزعة: {ctx['total_lubricants']:,} مزلقا."
        elif "مزلق" in ntx and "متوسط التوزيع" in ntx:
            new_text = f"متوسط التوزيع: {avg_lube:,} مزلقا لكل مستفيد ."
        elif "شهدت الفتره تقديم خدمات البرنامج" in ntx or "شهدت الفترة تقديم خدمات البرنامج" in tx:
            new_text = (
                f"شهدت الفترة تقديم خدمات البرنامج إلى {ctx['total_contacts']:,} مستفيدا، منهم "
                f"{ctx['base_visits']:,} مستفيدا جديدا و{ctx['followup_visits']:,} مستفيدا في المتابعة، "
                f"مما يعكس استمرار جهود الوصول إلى الفئات المستهدفة مع الحفاظ على متابعة المستفيدين السابقين."
            )
        elif "تم اجراء" in ntx and "اختبارا لفيروس نقص المناعه" in ntx:
            new_text = (
                f"تم إجراء {ctx['basic_tests']:,} اختبارا لفيروس نقص المناعة البشرية، مع تسجيل "
                f"{ctx['basic_negative']:,} نتيجة سلبية و{ctx['basic_positive']:,} نتيجة إيجابية."
            )
        elif "3 حالات دعم نفسي فردي" in tx or "حالات دعم نفسي فردي" in ntx:
            new_text = f"{ctx['psychosocial']:,} حالات دعم نفسي فردي ."

        if new_text is not None:
            _replace_in_paragraph(p, new_text)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def build_me_report_docx(ctx: Dict[str, Any]) -> bytes:
    """Fill the user's M&E report DOCX template and replace embedded chart images."""
    from docx import Document

    doc = Document(_template_path("Q2 2026 ME Report Final.docx"))
    label_en = str(ctx.get("period_label_en", ctx.get("period_label", "")))

    # Header/title replacements
    # Period words are dynamic. Q2 remains only when the selected period is exactly Q2.
    _pre_target_key = str(ctx.get("target_key") or _infer_target_key(ctx))
    _pre_scope = _period_scope(ctx, "annual" if _pre_target_key == "Annual" else _pre_target_key)
    _scope_title = "Q2 2026 Quarterly Report" if _pre_scope["short"] == "Q2" else f"{_pre_scope['short']} M&E Report"
    _rapid_total_label = f"Rapid test - positive ({_pre_scope['short']} total)" if _pre_scope["short"].startswith("Q") else f"Rapid test - positive ({_pre_scope['actual']} total)"
    _counted_result_label = "Counted Result to End of Q2" if _pre_scope["short"] == "Q2" else f"Counted Result ({_pre_scope['actual']})"
    replacements = {
        "Q2 2026 Quarterly Report (April – June 2026)": f"M&E Report ({label_en})",
        "Q2 2026 Quarterly Report": _scope_title,
        "Report date: July 2026": f"Report date: {datetime.now().strftime('%B %Y')}",
        "April – June 2026": label_en,
        "Q2 2026": label_en,
        "Performance Against Q2 Targets": f"Performance Against {_pre_scope['performance']} Targets",
        "Q2 reach or HIV-testing targets": f"{_pre_scope['performance']} reach or HIV-testing targets",
        "Q2 reach or testing targets": f"{_pre_scope['performance']} reach or testing targets",
        "Actual\n(Q2)": f"Actual\n({_pre_scope['actual']})",
        "Actual (Q2)": f"Actual ({_pre_scope['actual']})",
        "Q2 Target": f"{_pre_scope['target']} Target" if _pre_scope["short"] != "Q2" else "Q2 Target",
        "Q2 Total": "Q2 Total" if _pre_scope["short"] == "Q2" else "Total",
        "Rapid test - positive (Q2 total)": _rapid_total_label,
        "Rapid test – positive (Q2 total)": _rapid_total_label,
        "Counted Result to End of\nQ2": _counted_result_label,
        "Counted Result to End of Q2": _counted_result_label,
        "during Q2:": "during Q2:" if _pre_scope["short"] == "Q2" else "during the selected period:",
        "during Q2": "during Q2" if _pre_scope["short"] == "Q2" else "during the selected period",
        "Q2 total": f"{_pre_scope['short']} total" if _pre_scope["short"].startswith("Q") else f"{_pre_scope['actual']} total",
        "Quarterly Report": "Quarterly Report" if _pre_scope["short"].startswith("Q") else "M&E Report",
    }
    _replace_text_everywhere(doc, replacements)

    # Performance against targets table (template table 0)
    _target_key = str(ctx.get("target_key") or _infer_target_key(ctx))
    if _target_key == "Annual":
        _target_key = "annual"
    _target_lbl = _target_label(ctx, _target_key)
    _reach_target = _get_target(ctx, "reached", _target_key)
    _tests_target = _get_target(ctx, "hiv_tests", _target_key)
    _positive_target = _get_target(ctx, "positive", _target_key)
    _referral_target = _get_target(ctx, "referrals", _target_key)
    _scope = _period_scope(ctx, _target_key)
    _target_lbl = _scope["target"]

    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        # Keep the official table style, but fill target and achievement from the UI.
        if len(t0.rows) > 0:
            if len(t0.rows[0].cells) > 1:
                _safe_set_cell(t0.rows[0].cells[1], f"Actual ({_scope['actual']})")
            if len(t0.rows[0].cells) > 2:
                _safe_set_cell(t0.rows[0].cells[2], f"{_target_lbl} Target")
        rows = [
            (1, [ctx["base_visits"], _target_or_dash(_reach_target), _achievement(ctx["base_visits"], _reach_target)]),
            (2, [ctx["basic_tests"], _target_or_dash(_tests_target), _achievement(ctx["basic_tests"], _tests_target)]),
            (3, [ctx["rapid_positive_total"], _target_or_dash(_positive_target), _achievement(ctx["rapid_positive_total"], _positive_target)]),
            (4, [ctx["confirm_positive"], _target_or_dash(_referral_target), _achievement(ctx["confirm_positive"], _referral_target)]),
            (5, [ctx["rapid_positive_unconfirmed"], "-", "N/A"]),
            (6, [f"{ctx['no_test_result']:,} ({_pct(ctx['no_test_result'], ctx['base_visits'])})", "-", "N/A"]),
        ]
        for r_idx, vals in rows:
            if r_idx < len(t0.rows):
                for i, val in enumerate(vals, start=1):
                    if i < len(t0.rows[r_idx].cells):
                        _safe_set_cell(t0.rows[r_idx].cells[i], val)

    # Follow-up note box (template table 1)
    if len(doc.tables) > 1 and doc.tables[1].rows and doc.tables[1].rows[0].cells:
        _safe_set_cell(
            doc.tables[1].rows[0].cells[0],
            f"Note on Follow-up Visits (not counted in target)\n"
            f"{ctx['followup_visits']:,} beneficiaries received documented follow-up visits during the selected period. "
            f"Follow-up activity delivered {ctx['followup_syringes']:,} syringes, {ctx['followup_condoms']:,} condoms "
            f"and {ctx['followup_lubricants']:,} lubricants, plus {ctx['followup_tests']:,} HIV rapid tests. "
            f"Follow-up visits and follow-up HIV tests are reported separately and excluded from target calculations."
        )

    # Cascade table (template table 2)
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        if len(t2.rows) > 6 and len(t2.rows[6].cells) > 0:
            _safe_set_cell(t2.rows[6].cells[0], _rapid_total_label)
        cascade_rows = [
            (1, [ctx["base_visits"], "-"]),
            (2, [ctx["basic_tests"], f"{_pct(ctx['basic_tests'], ctx['base_visits'])} of basic reach"]),
            (3, [ctx["basic_negative"], f"{_pct(ctx['basic_negative'], ctx['basic_tests'])} of basic tests"]),
            (4, [ctx["basic_positive"], f"{_pct(ctx['basic_positive'], ctx['basic_tests'])} of basic tests"]),
            (5, [ctx["followup_tests"], "Reported separately"]),
            (6, [ctx["rapid_positive_total"], "Basic + follow-up"]),
            (7, [ctx["confirm_positive"], f"{_pct(ctx['confirm_positive'], ctx['rapid_positive_total'])} of rapid-positive"]),
            (8, [ctx["rapid_positive_unconfirmed"], f"{_pct(ctx['rapid_positive_unconfirmed'], ctx['rapid_positive_total'])} of rapid-positive"]),
        ]
        for r_idx, vals in cascade_rows:
            if r_idx < len(t2.rows):
                for i, val in enumerate(vals, start=1):
                    if i < len(t2.rows[r_idx].cells):
                        _safe_set_cell(t2.rows[r_idx].cells[i], val)

    # Commodities table (template table 3)
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        commodity_rows = [
            (1, [ctx["base_syringes"], ctx["followup_syringes"], ctx["total_syringes"]]),
            (2, [ctx["base_condoms"], ctx["followup_condoms"], ctx["total_condoms"]]),
            (3, [ctx["base_lubricants"], ctx["followup_lubricants"], ctx["total_lubricants"]]),
        ]
        for r_idx, vals in commodity_rows:
            if r_idx < len(t3.rows):
                for i, val in enumerate(vals, start=1):
                    if i < len(t3.rows[r_idx].cells):
                        _safe_set_cell(t3.rows[r_idx].cells[i], val)

    # Final safety pass for table/header labels that can keep period-specific wording.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                ctext = cell.text
                if "Counted Result to End of" in ctext or "Actual (Q2)" in ctext or "Q2 Total" in ctext or "Q2 Target" in ctext:
                    new_ctext = ctext.replace("Counted Result to End of\nQ2", _counted_result_label)
                    new_ctext = new_ctext.replace("Counted Result to End of Q2", _counted_result_label)
                    new_ctext = new_ctext.replace("Actual (Q2)", f"Actual ({_scope['actual']})")
                    new_ctext = new_ctext.replace("Q2 Total", "Q2 Total" if _scope["short"] == "Q2" else "Total")
                    new_ctext = new_ctext.replace("Q2 Target", f"{_scope['target']} Target" if _scope["short"] != "Q2" else "Q2 Target")
                    if new_ctext != ctext:
                        _safe_set_cell(cell, new_ctext)

    # Update major narrative paragraphs.
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("Between") or txt.startswith("During the selected period") or txt.startswith("During "):
            _target_lbl_clean = str(_target_lbl).strip()
            _target_lbl_display = _target_lbl_clean if _target_lbl_clean.upper().startswith("Q") else _target_lbl_clean.lower()
            _article = "an" if _target_lbl_display.lower().startswith("annual") else "a"
            _reach_target_txt = (
                f" against {_article} {_target_lbl_display} target of {_reach_target:,} "
                f"({_achievement(ctx['base_visits'], _reach_target)} achievement)"
                if _reach_target else ""
            )
            _tests_target_txt = (
                f" against {_article} {_target_lbl_display} target of {_tests_target:,} "
                f"({_achievement(ctx['basic_tests'], _tests_target)} achievement)"
                if _tests_target else ""
            )
            _replace_in_paragraph(
                p,
                f"During {label_en}, the Befrienders program reached {ctx['base_visits']:,} beneficiaries through basic visits"
                f"{_reach_target_txt} and conducted {ctx['basic_tests']:,} HIV rapid tests during basic visits"
                f"{_tests_target_txt}. Follow-up visits ({ctx['followup_visits']:,}) and follow-up HIV tests "
                f"({ctx['followup_tests']:,}) are reported separately."
            )
        elif "Beneficiaries reached:" in txt:
            _replace_in_paragraph(p, f"Beneficiaries reached: {ctx['base_visits']:,}" + (f" (target {_reach_target:,}) - {_achievement(ctx['base_visits'], _reach_target)}" if _reach_target else ""))
        elif "HIV rapid tests conducted" in txt:
            _replace_in_paragraph(p, f"HIV rapid tests conducted during basic visits: {ctx['basic_tests']:,}" + (f" (target {_tests_target:,}) - {_achievement(ctx['basic_tests'], _tests_target)}" if _tests_target else ""))
        elif "Rapid-test positive:" in txt:
            _replace_in_paragraph(p, f"Rapid-test positive: {ctx['rapid_positive_total']:,} total ({ctx['basic_positive']:,} during basic visits and {ctx['followup_positive']:,} during follow-up)")
        elif "Confirmed, referred and linked" in txt:
            _replace_in_paragraph(p, f"Confirmed, referred and linked to treatment: {ctx['confirm_positive']:,}; rapid-positive not confirmed: {ctx['rapid_positive_unconfirmed']:,}")
        elif "No documented HIV rapid test/result" in txt:
            _replace_in_paragraph(p, f"No documented HIV rapid test/result during basic visits: {ctx['no_test_result']:,} beneficiaries ({_pct(ctx['no_test_result'], ctx['base_visits'])})")
        elif txt.startswith("Follow-up visits to previously enrolled"):
            _replace_in_paragraph(p, f"Follow-up visits to previously enrolled beneficiaries are not counted against the {_scope['performance']} reach or HIV-testing targets. They are reported separately below and in Section 6 as a continuity-of-care indicator.")
        elif txt.startswith("Outreach activity"):
            _replace_in_paragraph(p, f"Outreach activity for the selected period reached {ctx['base_visits']:,} beneficiaries through basic visits.")
        elif txt.startswith("April:") or (" | " in txt and re.search(r"20\d{2}-\d{2}:", txt)):
            trend_txt = _monthly_trend_text(ctx)
            if trend_txt:
                _replace_in_paragraph(p, trend_txt)
        elif txt.startswith("The vast majority") or txt.startswith("Beneficiaries reached were male"):
            _replace_in_paragraph(p, f"Beneficiaries reached were male ({ctx['male']:,}, {_pct(ctx['male'], ctx['base_visits'])}), with {ctx['female']:,} female beneficiaries ({_pct(ctx['female'], ctx['base_visits'])}).")
        elif txt.startswith("The 20–44 age band") or txt.startswith("The 20-44 age band"):
            _replace_in_paragraph(p, _age_narrative(ctx))
        elif txt.startswith("After consolidating"):
            top = ctx.get("geo_counts")
            if isinstance(top, pd.DataFrame) and not top.empty:
                top_txt = ", ".join([f"{r['المنطقة']} ({int(r['العدد']):,})" for _, r in top.head(6).iterrows()])
            else:
                top_txt = "no area data available"
            _replace_in_paragraph(p, f"After consolidating area-name variants, geographic distribution for the selected period was: {top_txt}.")
        elif txt.startswith("Basic-visit testing and follow-up testing"):
            _replace_in_paragraph(p, f"Basic-visit testing and follow-up testing are presented separately to avoid mixing different denominators. Confirmed and linked cases: {ctx['confirm_positive']:,}; unconfirmed rapid-positive results: {ctx['rapid_positive_unconfirmed']:,}.")
        elif txt.startswith("Across basic and follow-up"):
            _replace_in_paragraph(p, f"Across basic and follow-up activity, {ctx['total_tests']:,} HIV rapid tests were documented: {ctx['total_negative']:,} negative and {ctx['rapid_positive_total']:,} rapid-positive. Confirmed and linked cases: {ctx['confirm_positive']:,}; unconfirmed rapid-positive results: {ctx['rapid_positive_unconfirmed']:,}.")
        elif txt.startswith("Syringes remained"):
            _replace_in_paragraph(p, f"Syringes remained the dominant prevention commodity. The dataset records {ctx['base_syringes']:,} syringes during basic visits and {ctx['followup_syringes']:,} during follow-up visits, for a total of {ctx['total_syringes']:,}.")
        elif txt.startswith("Beyond HIV testing and commodity distribution"):
            _replace_in_paragraph(p, f"Beyond HIV testing and commodity distribution, the program delivered the following complementary services across basic and follow-up visits during {('Q2' if _scope['short'] == 'Q2' else 'the selected period')}:" )
        elif "basic-visit records" in txt and "no documented rapid-test result" in txt:
            _replace_in_paragraph(p, f"{ctx['no_test_result']:,} basic-visit records ({_pct(ctx['no_test_result'], ctx['base_visits'])}) have no documented rapid-test result. The dataset confirms that no test or result was recorded, but it does not document the reason. These records should not be classified as test refusals without a specific reason field.")
        elif txt.startswith("Both headline Q2 targets") or txt.startswith("Both headline"):
            if _reach_target and _tests_target:
                _replace_in_paragraph(p, f"Headline target performance for the selected period: beneficiary reach achieved {_achievement(ctx['base_visits'], _reach_target)} and basic-visit HIV testing achieved {_achievement(ctx['basic_tests'], _tests_target)}.")
            else:
                _replace_in_paragraph(p, "Headline target performance is presented in Section 2 where target values are available.")
        elif txt.startswith("Continue reporting follow-up activity"):
            _replace_in_paragraph(p, f"Continue reporting follow-up activity ({ctx['followup_visits']:,} beneficiaries) as a separate continuity-of-care indicator and do not merge it into primary reach or testing-target calculations.")
        elif txt.startswith("Review the monthly pattern"):
            trend_txt = _monthly_trend_text(ctx)
            if trend_txt:
                _replace_in_paragraph(p, f"Review the monthly pattern before planning the next period: {trend_txt}. Investigate staffing, site access and operational scheduling where performance changes materially.")
        elif "Before submission" in txt and "144-condom" in txt:
            _replace_in_paragraph(p, "Before submission, verify any high-value commodity records against the source form and rerun the DataBridge quality assessment on the revised dataset.")
        elif "Syphilis testing:" in txt:
            _replace_in_paragraph(p, f"Syphilis testing: {ctx['syphilis']:,} beneficiaries")
        elif "Psychosocial support" in txt:
            _replace_in_paragraph(p, f"Psychosocial support sessions: {ctx['psychosocial']:,} beneficiaries")
        elif "OST / methadone" in txt:
            _replace_in_paragraph(p, f"OST / methadone referrals: {ctx['methadone_referrals']:,} beneficiaries")

    # Clean Word metadata so exported files do not retain old template titles/subjects.
    try:
        props = doc.core_properties
        props.title = f"Befrienders Program — {_scope['performance']} M&E Report"
        props.subject = f"M&E Report ({label_en})"
        props.keywords = "Befrienders, M&E, UNDP, Global Fund, HIV"
        props.comments = "Generated by DataBridge Pro"
        props.author = "DataBridge Pro"
        props.last_modified_by = "DataBridge Pro"
        props.category = "M&E Report"
    except Exception:
        pass

    bio = io.BytesIO(); doc.save(bio)
    docx_bytes = bio.getvalue()
    image_map = {
        "word/media/image1.png": _matplotlib_chart_png("monthly", ctx, 1050, 570),
        "word/media/image2.png": _matplotlib_chart_png("gender", ctx, 750, 600),
        "word/media/image3.png": _matplotlib_chart_png("age", ctx, 1050, 570),
        "word/media/image4.png": _matplotlib_chart_png("geo", ctx, 1115, 590),
        "word/media/image5.png": _matplotlib_chart_png("commodities", ctx, 1039, 590),
    }
    docx_bytes = _replace_docx_media(docx_bytes, image_map)
    return _clean_docx_extended_properties(
        docx_bytes,
        title=f"Befrienders Program — {_scope['performance']} M&E Report",
        subject=f"M&E Report ({label_en})",
    )




def _age_tool_rows(ctx: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Return age-band commodity totals from base visits for official narrative use."""
    df = ctx.get("base_df")
    if not isinstance(df, pd.DataFrame) or df.empty:
        return []
    age_col = find_col(df, ["السن"]) or find_col(df, ["العمر"]) or find_col(df, ["age"])
    if not age_col:
        return []
    syr_col = find_best_quantity_col(df, "سرنجات")
    cond_col = find_best_quantity_col(df, "واقيات")
    lube_col = find_best_quantity_col(df, "مزلقات")
    order = ['15 - 19', '20 - 24', '25 - 29', '30 - 34', '35 - 39', '40 - 44', '45 - 49', '50 او اكثر', '50+']
    rows = []
    for age in order:
        part = df[df[age_col].astype(str).str.strip().eq(age)]
        if part.empty:
            continue
        rows.append({
            "السن": age,
            "واقيات": int(pd.to_numeric(part[cond_col], errors="coerce").fillna(0).sum()) if cond_col else 0,
            "مزلقات": int(pd.to_numeric(part[lube_col], errors="coerce").fillna(0).sum()) if lube_col else 0,
            "سرنجات": int(pd.to_numeric(part[syr_col], errors="coerce").fillna(0).sum()) if syr_col else 0,
        })
    rows.sort(key=lambda r: r["سرنجات"], reverse=True)
    return rows

def build_undp_quarterly_arabic_docx(ctx: Dict[str, Any]) -> bytes:
    """Fill the official UNDP Arabic quarterly form in-place.

    This is a strict template-fill path: it opens
    templates/تقرير الربع الثاني 2026.docx and updates only existing table cells
    and value paragraphs. It does not rebuild the form, remove tables, or change
    the official layout.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(_template_path("تقرير الربع الثاني 2026.docx"))
    period_doc = _period_doc_label(ctx)

    def set_cell(cell, value: Any) -> None:
        _safe_set_cell(cell, _fmt_int(value) if isinstance(value, int) else value)

    def set_row(row, cell_index: int, value: Any) -> None:
        if 0 <= cell_index < len(row.cells):
            set_cell(row.cells[cell_index], value)

    # General period in first information table.
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        for row in t0.rows:
            txt = " ".join(c.text for c in row.cells)
            if "فترة التقرير" in txt or "Periodicity" in txt:
                if len(row.cells) > 1:
                    set_cell(row.cells[1], period_doc)

    # Main PWID quantitative section. The table layout is fixed in the official
    # form, so use row numbers from the template and preserve all other cells.
    if len(doc.tables) > 1:
        t = doc.tables[1]
        # Main indicators: result and achievement columns.
        if len(t.rows) > 4:
            reach_target = num(t.rows[3].cells[1].text) or 0
            test_target = num(t.rows[4].cells[1].text) or 0
            set_row(t.rows[3], 2, ctx["base_visits"])
            set_row(t.rows[3], 3, _pct(ctx["base_visits"], reach_target) if reach_target else t.rows[3].cells[3].text)
            set_row(t.rows[4], 2, ctx["basic_tests"])
            set_row(t.rows[4], 3, _pct(ctx["basic_tests"], test_target) if test_target else t.rows[4].cells[3].text)
        # Positivity.
        if len(t.rows) > 9:
            set_row(t.rows[7], 1, ctx["rapid_positive_total"])
            set_row(t.rows[7], 2, f"نسبة {_pct(ctx['rapid_positive_total'], max(ctx['total_contacts'], ctx['base_visits']))} حالات ايجابي من اجمالي المستفيدين")
            set_row(t.rows[8], 1, ctx["confirm_positive"])
            set_row(t.rows[9], 1, ctx["referrals_linked"])
        # Prevention package: basic visits only, per official instruction in the form.
        if len(t.rows) > 18:
            avg_syr = round(ctx["base_syringes"] / max(ctx["base_visits"], 1))
            avg_cond = ctx["base_condoms"] / max(ctx["base_visits"], 1)
            avg_lube = round(ctx["base_lubricants"] / max(ctx["base_visits"], 1))
            set_row(t.rows[12], 1, ctx["base_syringes"])
            set_row(t.rows[12], 2, f"متوسط التوزيع {avg_syr} سرنجه للفرد")
            set_row(t.rows[13], 1, ctx["base_condoms"])
            set_row(t.rows[13], 2, f"متوسط التوزيع {avg_cond:.1f} واقي للفرد")
            set_row(t.rows[14], 1, ctx["base_lubricants"])
            set_row(t.rows[14], 2, f"متوسط توزيع {avg_lube} مزلق للفرد")
            set_row(t.rows[15], 1, ctx["base_visits"])
            set_row(t.rows[15], 2, "اجمالي المستفيدين تلقو رسائل توعيه لفيروس نقص المناعه ")
            set_row(t.rows[16], 1, ctx["psychosocial"])
            set_row(t.rows[17], 1, ctx["syphilis"])
            set_row(t.rows[18], 1, ctx["methadone_referrals"])
        # Retention.
        if len(t.rows) > 24:
            set_row(t.rows[21], 1, ctx["followup_visits"])
            set_row(t.rows[22], 1, ctx["followup_syringes"])
            set_row(t.rows[23], 1, ctx["followup_condoms"])
            set_row(t.rows[24], 1, ctx["followup_lubricants"])
        # Female PWID block.
        if len(t.rows) > 31:
            set_row(t.rows[27], 1, ctx["female"])
            set_row(t.rows[28], 1, ctx.get("basic_tests_female", 0))
            set_row(t.rows[29], 1, ctx.get("rapid_positive_total_female", 0))
            set_row(t.rows[30], 1, ctx.get("confirm_positive_female", 0))
            # Linkage is usually the same as confirmed among females unless a
            # dedicated gender referral field is available.
            set_row(t.rows[31], 1, ctx.get("confirm_positive_female", 0))

    # Section B narrative: update values only in existing paragraphs.
    age_tools = _age_tool_rows(ctx)
    top_age = age_tools[0] if age_tools else None
    geo_counts = ctx.get("geo_counts")
    if isinstance(geo_counts, pd.DataFrame) and not geo_counts.empty:
        geo_top = geo_counts.head(4)
        geo_top_lines = [f"{r['المنطقة']} ({int(r['العدد']):,})" for _, r in geo_top.iterrows()]
    else:
        geo_top_lines = []

    for p in doc.paragraphs:
        tx = p.text.strip()
        if not tx:
            continue
        ntx = norm_ar(tx)
        new_text = None
        if "اجمالي عدد المستفيدين" in ntx:
            new_text = f"إجمالي عدد المستفيدين : {ctx['total_contacts']:,}"
        elif "اجمالي الوصول زيارات اول مره" in ntx:
            new_text = f"إجمالي الوصول زيارات أول مرة: {ctx['base_visits']:,}"
        elif "اجمالي زيارات المتابعه" in ntx:
            new_text = f"إجمالي زيارات المتابعة: {ctx['followup_visits']:,}"
        elif "اجمالى التحاليل" in ntx and "اول مره" not in ntx and "متابعه" not in ntx:
            new_text = f"اجمالى التحاليل : {ctx['basic_tests']:,}"
        elif "اجمالي التحاليل اول مره" in ntx:
            new_text = f"إجمالي التحاليل اول مره : {ctx['basic_tests']:,} تحليل"
        elif "اجمالى تحاليل متابعه" in ntx:
            new_text = f"اجمالى تحاليل متابعه : {ctx['followup_tests']:,} تحليل"
        elif "اختبارات الزهري" in ntx:
            new_text = f"اختبارات الزهري: {ctx['syphilis']:,}"
        elif "الدعم النفسي" in ntx and "فردي" in ntx:
            new_text = f"الدعم النفسي فردي : {ctx['psychosocial']:,}"
        elif "عدد الحالات الايجابيه" in ntx:
            new_text = f"عدد الحالات الإيجابية: {ctx['rapid_positive_total']:,} حاالت"
        elif "عدد الحالات المؤكده" in ntx:
            new_text = f"عدد الحالات المؤكدة : {ctx['confirm_positive']:,} حاالت"
        elif "عدد الاحالات لصرف العلاج" in ntx:
            new_text = f"عدد الإحالات لصرف العلاج : {ctx['referrals_linked']:,} حاالت"
        elif "احالات طبيه اخري" in ntx or "احالات طبية اخرى" in ntx:
            new_text = f"إحالات طبية أخرى: {ctx['medical_referrals']:,}"
        elif "احالات الميثادون" in ntx:
            new_text = f"إحالات الميثادون : {ctx['methadone_referrals']:,}"
        elif "اجمالي" in ntx and "تحليل خلال الربع" in ntx:
            new_text = f"إجمالي {ctx['basic_tests']:,} تحليل خلال الربع"
        elif "تم اكتشاف" in ntx and "حالات ايجابيه" in ntx:
            new_text = f"تم اكتشاف {ctx['rapid_positive_total']:,} حالات إيجابية ({ctx['confirm_positive']:,} مؤكدة)"
        elif "حالتان" in ntx and "رفض سحب العينه" in ntx:
            new_text = f"{ctx['rapid_positive_unconfirmed']:,} حالة لم تستكمل التحليل التأكيدي لرفض سحب العينة ."
        elif "احالات العلاج" in ntx:
            new_text = f"إحالات العلاج : {ctx['referrals_linked']:,} فقط"
        elif "تقديم" in ntx and "جلسات خلال الربع" in ntx:
            new_text = f"تقديم {ctx['psychosocial']:,} جلسات خلال الربع"
        elif "سرنجات:" in tx and top_age:
            new_text = f"سرنجات: {int(top_age['سرنجات']):,}"
        elif "مزلقات:" in tx and top_age:
            new_text = f"مزلقات: {int(top_age['مزلقات']):,}"
        elif "واقيات:" in tx and top_age:
            new_text = f"واقيات: {int(top_age['واقيات']):,}"
        elif "اجمالي الارقام" in ntx:
            new_text = f"إجمالي الأرقام تقريبا : {ctx['base_visits']:,}"
        elif "اعلي المناطق" in ntx and geo_top_lines:
            # Keep header text as-is.
            new_text = tx
        elif "اكتوبر" in ntx and "اعلي منطقه" in ntx and geo_top_lines:
            new_text = f"{geo_top_lines[0]} أعلى منطقة"

        if new_text is not None:
            _replace_in_paragraph(p, new_text)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Age/protection table in Section B: it is an embedded table-like docx table
    # only in the original visual form. In this template it appears as actual text
    # paragraphs around a chart image, so we do not attempt to rebuild it here.
    # This preserves the official form. Dynamic age tables are available in the
    # separate M&E report and Summary Excel.

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def build_donor_package(ctx: Dict[str, Any], package_type: str) -> bytes:
    """Return a ZIP file containing template-based reports for the selected package."""
    label = re.sub(r"[^0-9A-Za-z\u0600-\u06FF_-]+", "_", str(ctx.get("period_label", "period")))
    summary_xlsx = build_summary_excel_bytes(ctx)
    me_docx = build_me_report_docx(ctx)
    monthly_docx = build_monthly_arabic_report_docx(ctx)
    undp_docx = build_undp_quarterly_arabic_docx(ctx)

    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as z:
        if package_type == "single_month":
            z.writestr(f"Monthly_Report_{label}.docx", monthly_docx)
            z.writestr(f"ME_Report_{label}.docx", me_docx)
            z.writestr(f"Summary_{label}.xlsx", summary_xlsx)
        elif package_type == "range":
            z.writestr(f"ME_Report_{label}.docx", me_docx)
            z.writestr(f"Summary_{label}.xlsx", summary_xlsx)
        elif package_type == "undp_quarterly":
            z.writestr(f"UNDP_Quarterly_Report_{label}.docx", undp_docx)
        else:
            z.writestr(f"ME_Report_{label}.docx", me_docx)
            z.writestr(f"Summary_{label}.xlsx", summary_xlsx)
    zbuf.seek(0)
    return zbuf.getvalue()
