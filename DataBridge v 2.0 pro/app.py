"""
DataBridge — M&E Hub for Health Programs
Complete Monitoring & Evaluation Platform
Phase 1 — Standalone HIV/IDUs Data Analysis App
"""

import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
import io
import os
import json
import logging
import html
from datetime import datetime
from collections import deque
from typing import Optional, List, Dict, Any, Tuple
import openpyxl
from io import BytesIO

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# ────────────────────────────────────────────────────────────────
# CONSTANTS
# ────────────────────────────────────────────────────────────────
MAX_HISTORY        = 10
MAX_UPLOAD_SIZE_MB = 100
MAX_HISTORY_MEM_MB = 500
USERS_FILE         = "users.json"

# ────────────────────────────────────────────────────────────────
# MODULAR IMPORTS — DataBridge v2.0
# ────────────────────────────────────────────────────────────────
from modules.settings import T, APP_VERSION, load_config, save_config, get_appdata_dir
from modules.utils import safe_html, validate_uploaded_file, safe_read_excel, SmartHistoryManager, secure_multi_condition_filter
from modules.auth import render_login, change_password
from modules.quality_engine import run_quality_engine, compute_quality_score
from modules.source_converter import format_idus_source_date
from modules.data_cleaner import smart_read_excel, clean_dataframe, report_to_dataframe, mapping_report_to_dataframe, generate_standard_df, get_mapping_summary, save_mapping_memory_from_report
from modules.data_repair_center import render_data_repair_center

# ────────────────────────────────────────────────────────────────
# TRANSLATIONS
# ────────────────────────────────────────────────────────────────
# ────────────────────────────────────────────────────────────────
# ────────────────────────────────────────────────────────────────
#  PAGE CONFIG
# ════════════════════════════════════════
st.set_page_config(
    page_title="DataBridge — M&E Hub",
    layout="wide",
    page_icon="🌉",
    initial_sidebar_state="expanded"
)

# ════════════════════════════════════════
#  GLOBAL STYLES
# ════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

html, body, [class*="css"] { font-family: 'Space Grotesk', sans-serif; }
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 0 2rem 3rem 2rem !important; max-width: 100% !important; }

section[data-testid="stSidebar"] {
    background: #0a0a0f;
    border-right: 1px solid #1e1e2e;
}
section[data-testid="stSidebar"] * { color: #e0e0f0 !important; }

.top-header {
    background: linear-gradient(135deg, #0a0a0f 0%, #0d0d1a 100%);
    border-bottom: 1px solid #1e1e2e;
    padding: 1.2rem 2rem;
    margin: -1rem -2rem 2rem -2rem;
    display: flex;
    align-items: center;
    gap: 1rem;
}
.top-header .logo { font-size: 1.5rem; font-weight: 700; color: #fff; letter-spacing: -0.5px; }
.top-header .logo span { color: #7c6aff; }
.top-header .subtitle { font-size: 0.78rem; color: #666; font-weight: 400; }
.header-badge {
    margin-left: auto;
    background: #7c6aff22;
    border: 1px solid #7c6aff44;
    color: #a89fff;
    padding: 0.25rem 0.75rem;
    border-radius: 20px;
    font-size: 0.72rem;
    font-weight: 500;
}
.section-header {
    display: flex;
    align-items: center;
    gap: 0.6rem;
    margin: 1.8rem 0 1rem 0;
    padding-bottom: 0.6rem;
    border-bottom: 1px solid #1e1e2e;
}
.section-header .icon {
    width: 32px; height: 32px;
    background: #7c6aff22;
    border: 1px solid #7c6aff44;
    border-radius: 8px;
    display: flex; align-items: center; justify-content: center;
    font-size: 0.9rem;
}
.section-header h2 { font-size: 1rem; font-weight: 600; color: #e0e0f0; margin: 0; }

.metrics-row { display: grid; grid-template-columns: repeat(5, 1fr); gap: 12px; margin-bottom: 1.5rem; }
.metric-card {
    background: #0d0d1a;
    border: 1px solid #1e1e2e;
    border-radius: 12px;
    padding: 1rem 1.2rem;
    transition: border-color 0.2s;
}
.metric-card:hover { border-color: #7c6aff44; }
.metric-card .label { font-size: 0.72rem; color: #555; text-transform: uppercase; letter-spacing: 0.08em; margin-bottom: 0.4rem; }
.metric-card .value { font-size: 1.6rem; font-weight: 700; color: #fff; font-family: 'JetBrains Mono', monospace; line-height: 1; }
.metric-card .sub { font-size: 0.72rem; color: #7c6aff; margin-top: 0.3rem; }

.stButton > button {
    background: #7c6aff !important;
    color: #fff !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Space Grotesk', sans-serif !important;
    font-weight: 500 !important;
    font-size: 0.85rem !important;
    padding: 0.45rem 1.2rem !important;
    transition: all 0.2s !important;
}
.stButton > button:hover { background: #6855e8 !important; transform: translateY(-1px); }

.stDataFrame { border-radius: 10px !important; overflow: hidden; }
[data-testid="stDataFrame"] { border: 1px solid #1e1e2e !important; border-radius: 10px; }

.stSelectbox > div > div, .stTextInput > div > div > input {
    background: #0d0d1a !important;
    border: 1px solid #2a2a4e !important;
    border-radius: 8px !important;
    color: #e0e0f0 !important;
    font-family: 'Space Grotesk', sans-serif !important;
}

.stSuccess { background: #0a1f0a !important; border: 1px solid #1a4a1a !important; border-radius: 8px !important; }
.stWarning { background: #1f1a0a !important; border: 1px solid #4a3a0a !important; border-radius: 8px !important; }
.stError   { background: #1f0a0a !important; border: 1px solid #4a1a1a !important; border-radius: 8px !important; }

/* Fixed sidebar: keep it visible, then hide Streamlit native collapse/toggle controls */
section[data-testid="stSidebar"] {
    display: block !important;
    visibility: visible !important;
    min-width: 18rem !important;
    width: 18rem !important;
    max-width: 18rem !important;
    transform: translateX(0) !important;
}

section[data-testid="stSidebar"] > div {
    display: block !important;
    visibility: visible !important;
}

button[data-testid="collapsedControl"],
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"],
[data-testid="stSidebarCollapseButton"] button {
    display: none !important;
    visibility: hidden !important;
    opacity: 0 !important;
    pointer-events: none !important;
}

.info-box {
    background: #0d0d1a;
    border: 1px solid #1e1e2e;
    border-left: 3px solid #7c6aff;
    border-radius: 8px;
    padding: 0.8rem 1rem;
    font-size: 0.82rem;
    color: #888;
    margin: 0.8rem 0;
}

.stTabs [data-baseweb="tab-list"] { gap: 4px; background: transparent; border-bottom: 1px solid #1e1e2e; }
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    border: none !important;
    border-bottom: 2px solid transparent !important;
    color: #555 !important;
    font-family: 'Space Grotesk', sans-serif !important;
    font-size: 0.85rem !important;
    padding: 0.5rem 1rem !important;
    border-radius: 0 !important;
}
.stTabs [aria-selected="true"] { color: #7c6aff !important; border-bottom-color: #7c6aff !important; }

hr { border-color: #1e1e2e !important; margin: 2rem 0 !important; }

[data-testid="stDownloadButton"] button {
    background: #1a1a2e !important;
    border: 1px solid #7c6aff44 !important;
    color: #a89fff !important;
}
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════
#  SESSION STATE INIT
# ════════════════════════════════════════
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if "lang" not in st.session_state:
    st.session_state["lang"] = "ar"
if "hdf" not in st.session_state:
    st.session_state["hdf"] = None
if "hiv_file_name" not in st.session_state:
    st.session_state["hiv_file_name"] = None
if "dq_notes" not in st.session_state:
    st.session_state["dq_notes"] = {}
if "dq_issues" not in st.session_state:
    st.session_state["dq_issues"] = None
if "targets" not in st.session_state:
    st.session_state["targets"] = {
        "hiv_tests": 0, "positive": 0, "referrals": 0,
        "followups": 0, "refusal": 0
    }
if "show_change_pw" not in st.session_state:
    st.session_state["show_change_pw"] = False
if "data_clean_report" not in st.session_state:
    st.session_state["data_clean_report"] = None
if "standard_df" not in st.session_state:
    st.session_state["standard_df"] = None
if "mapper_approved" not in st.session_state:
    st.session_state["mapper_approved"] = False


# ════════════════════════════════════════
#  LANGUAGE HELPER
# ════════════════════════════════════════
def t(key: str) -> str:
    return T[st.session_state["lang"]].get(key, key)


#  LOGIN GATE
# ════════════════════════════════════════
if not st.session_state["authenticated"]:
    # Language toggle on login page (top right)
    _ll1, _ll2 = st.columns([5, 1])
    with _ll2:
        if st.button(t("lang_btn"), key="lang_login", use_container_width=True):
            st.session_state["lang"] = "en" if st.session_state["lang"] == "ar" else "ar"
            st.rerun()
    render_login(st.session_state["lang"])
    st.stop()


# ── Force sidebar open even if browser remembered it collapsed ──
components.html("""
<script>
(function () {
  function openSidebarIfCollapsed() {
    const doc = window.parent.document;
    const sidebar = doc.querySelector('section[data-testid="stSidebar"]');
    const collapsedControl = doc.querySelector('[data-testid="collapsedControl"], button[data-testid="collapsedControl"]');
    let collapsed = false;
    if (sidebar) {
      const rect = sidebar.getBoundingClientRect();
      const style = window.parent.getComputedStyle(sidebar);
      collapsed = rect.width < 100 || rect.left < -50 || style.visibility === 'hidden';
    } else {
      collapsed = true;
    }
    if (collapsedControl && collapsed) {
      collapsedControl.click();
    }
  }
  openSidebarIfCollapsed();
  setTimeout(openSidebarIfCollapsed, 250);
  setTimeout(openSidebarIfCollapsed, 800);
  setTimeout(openSidebarIfCollapsed, 1500);
})();
</script>
""", height=0)

# ════════════════════════════════════════
#  HEADER (after login)
# ════════════════════════════════════════
st.markdown(f"""
<div class="top-header">
  <div>
    <div class="logo">Data<span>Bridge</span></div>
    <div class="subtitle">{t('app_sub')}</div>
  </div>
  <div class="header-badge">🌉 {APP_VERSION}</div>
</div>
""", unsafe_allow_html=True)

# ════════════════════════════════════════
#  SIDEBAR
# ════════════════════════════════════════
with st.sidebar:
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align:center;margin-bottom:1rem;">
      <div style="font-size:1.3rem;font-weight:700;color:#fff;">Data<span style="color:#7c6aff">Bridge</span></div>
      <div style="font-size:0.68rem;color:#444;">M&E Hub for Health Programs</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    # Language toggle
    if st.button(f"🌐 {t('lang_btn')}", use_container_width=True, key="lang_sidebar"):
        st.session_state["lang"] = "en" if st.session_state["lang"] == "ar" else "ar"
        st.rerun()

    st.markdown("---")

    # File info
    if st.session_state["hdf"] is not None:
        hdf = st.session_state["hdf"]
        st.markdown(f"""
        <div style="font-size:0.75rem;color:#555;padding:0.5rem 0;">
        📄 <b style="color:#888">{st.session_state['hiv_file_name']}</b><br>
        <span style="font-family:'JetBrains Mono',monospace;color:#7c6aff">{hdf.shape[0]:,}</span> rows &nbsp;·&nbsp;
        <span style="font-family:'JetBrains Mono',monospace;color:#7c6aff">{hdf.shape[1]}</span> cols
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # Logout
    st.markdown(f"<div style='font-size:0.75rem;color:#555;'>👤 {st.session_state.get('current_user','')}</div>", unsafe_allow_html=True)

    if st.button(t("change_password_btn"), use_container_width=True, key="toggle_pw"):
        st.session_state["show_change_pw"] = not st.session_state["show_change_pw"]

    if st.session_state["show_change_pw"]:
        with st.form("change_pw_form"):
            old_pw  = st.text_input(t("current_password"), type="password", key="old_pw")
            new_pw  = st.text_input(t("new_password"), type="password", key="new_pw")
            new_pw2 = st.text_input(t("confirm_password"), type="password", key="new_pw2")
            submitted = st.form_submit_button(t("save_btn"), use_container_width=True)
            if submitted:
                if new_pw != new_pw2:
                    st.error(t("password_mismatch"))
                else:
                    ok, msg_key = change_password(
                        st.session_state.get("current_user", "admin"),
                        old_pw, new_pw
                    )
                    if ok:
                        st.success(t(msg_key))
                        st.session_state["show_change_pw"] = False
                    else:
                        st.error(t(msg_key))

    st.markdown("---")
    if st.button(f"🚪 {t('logout')}", use_container_width=True, key="logout_btn"):
        for key in ["authenticated", "current_user", "hdf", "hiv_file_name", "dq_issues", "dq_notes", "data_clean_report", "standard_df"]:
            st.session_state.pop(key, None)
        st.rerun()

    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("<div style='font-size:0.68rem;color:#333;'>DataBridge © 2026</div>", unsafe_allow_html=True)



# ════════════════════════════════════════
#  IDUs SOURCE AUTO-CONVERTER HELPERS
# ════════════════════════════════════════
def _looks_like_idus_204_source(file_bytes: bytes) -> bool:
    """Detect the original 204-column IDUs source workbook before normal Smart Mapping."""
    try:
        sample = pd.read_excel(io.BytesIO(file_bytes), header=None, skiprows=4, nrows=10)
        if sample.shape[1] < 180:
            return False
        has_visit_date = 2 in sample.columns and pd.to_datetime(sample[2], errors="coerce").notna().sum() >= 1
        has_code = 15 in sample.columns and sample[15].notna().sum() >= 1
        return bool(has_visit_date and has_code)
    except Exception:
        return False


def _convert_idus_204_source_bytes(file_bytes: bytes) -> Tuple[pd.DataFrame, Any]:
    """Convert the original 204-column IDUs source workbook to the program format.
    Important: do not swap day/month for Excel datetime cells; that was the
    cause of wrong month totals in some versions.
    """
    import warnings as _ww
    _ww.filterwarnings("ignore")

    _dfs = pd.read_excel(io.BytesIO(file_bytes), header=None, skiprows=4)
    _dfs.columns = range(len(_dfs.columns))

    if _dfs.shape[1] < 180:
        raise ValueError("هذا الملف لا يبدو الشيت الأساسي 204 عمود.")

    _last_code = _dfs[15].last_valid_index() if 15 in _dfs.columns else None
    _last_date = _dfs[2].last_valid_index() if 2 in _dfs.columns else None
    _valid_last = [x for x in [_last_code, _last_date] if x is not None]
    if not _valid_last:
        raise ValueError("لم يتم العثور على تاريخ الزيارة أو الكود المجمع داخل الشيت الأساسي.")
    _last_row = max(_valid_last)
    _dfs = _dfs.iloc[:_last_row + 1].copy()

    def _c(i):
        return _dfs[i-1] if (i-1) in _dfs.columns else pd.Series([None] * len(_dfs), index=_dfs.index)

    def _yes(value: Any) -> bool:
        s = str(value).strip().lower()
        return s in ["نعم", "ايوه", "أيوه", "yes", "y", "1", "true"]

    def _yq(yn, qty):
        r = []
        for a, b in zip(yn, qty):
            if _yes(a):
                try:
                    r.append(int(float(b)) if pd.notna(b) else None)
                except Exception:
                    r.append(None)
            else:
                r.append(None)
        return r

    def _ys(s):
        def _one(x):
            if pd.isna(x) or str(x).strip() in ["", "nan", "None"]:
                return None
            return "نعم" if _yes(x) else "لا"
        return s.apply(_one)

    _cv = pd.DataFrame()
    _cv['مسلسل']                                         = _c(1)
    _cv['الكود المجمع']                                  = _c(16)
    _cv['اسم الجمعية']                                   = _c(2)
    _cv['تاريخ الزيارة']                                 = format_idus_source_date(_c(3), swap_excel_dates=False)
    _cv['اسم مندوب الوصول ثلاثي\n(الباحث الميداني) ']  = _c(4)
    _cv['محافظة الوصول للمستفيد']                       = _c(5)
    _cv['منطقة الوصول للمستفيد ']                       = _c(6)
    _cv['مكان المقابلة']                                 = _c(17)
    _cv['السن']                                          = _c(18)
    _cv['النوع']                                         = _c(19)
    _cv['واقيات']                                        = _yq(_c(27), _c(28))
    _cv['مزلقات']                                        = _yq(_c(31), _c(32))
    _cv[' زهري']                                         = _ys(_c(34))
    _cv['دعم نفسي']                                      = _ys(_c(36))
    _cv['سرنجات']                                        = _yq(_c(39), _c(40))
    _cv['ميثادون']                                       = _ys(_c(42))
    _cv['نتيجة التحليل ']                                = _c(48)
    _cv['نتيجة التحليل التاكيدي']                       = _c(50)
    _cv['الاحالة علي العلاج']                           = _c(51)

    _cv['زيارة متابعة 1']        = format_idus_source_date(_c(57), swap_excel_dates=False)
    _cv['واقيات متابعة 1']       = _yq(_c(62), _c(63))
    _cv['مزلقات متابعة 1']       = _yq(_c(66), _c(67))
    _cv['زهري متابعة 1']         = _ys(_c(69))
    _cv['دعم نفسي متابعة 1']     = _ys(_c(71))
    _cv['سرنجات متابعة 1']       = _yq(_c(74), _c(75))
    _cv['نتيجة تحليل متابعة 1']  = _c(78)
    _cv['ميثادون متابعة 1']      = _ys(_c(80))

    _cv['زيارة متابعة 2']        = format_idus_source_date(_c(86), swap_excel_dates=False)
    _cv['واقيات متابعة 2']       = _yq(_c(91), _c(92))
    _cv['مزلقات متابعة 2']       = _yq(_c(95), _c(96))
    _cv['زهري متابعة 2']         = _ys(_c(98))
    _cv['دعم نفسي متابعة 2']     = _ys(_c(100))
    _cv['سرنجات متابعة 2']       = _yq(_c(103), _c(104))
    _cv['نتيجة تحليل متابعة 2']  = _c(107)

    _cv['زيارة متابعة 3']        = format_idus_source_date(_c(115), swap_excel_dates=False)
    _cv['واقيات متابعة 3']       = _yq(_c(120), _c(121))
    _cv['مزلقات متابعة 3']       = _yq(_c(124), _c(125))
    _cv['نتيجة تحليل متابعة 3']  = _c(136)

    _cv['زيارة متابعة 4']        = format_idus_source_date(_c(144), swap_excel_dates=False)
    _cv['نتيجة تحليل متابعة 4']  = _c(165)

    _cv['زيارة متابعة 5']        = format_idus_source_date(_c(173), swap_excel_dates=False)
    _cv['نتيجة تحليل متابعة 5']  = _c(194)

    _cv = _cv.dropna(how='all').reset_index(drop=True)
    _cv.columns = [str(c).strip() for c in _cv.columns]
    _cv, _clean_report = clean_dataframe(_cv)
    return _cv, _clean_report


# ════════════════════════════════════════
#  MAIN CONTENT
# ════════════════════════════════════════
st.markdown(f'<div class="section-header"><div class="icon">🏥</div><h2>{t("section_title")}</h2></div>', unsafe_allow_html=True)

# ── Upload Mode Selector ──
_upload_mode = st.radio(
    "نوع الملف:" if st.session_state["lang"] == "ar" else "File type:",
    ["📊 شيت البرنامج (27 عمود)" if st.session_state["lang"] == "ar" else "📊 Program Sheet (27 cols)",
     "🔄 الشيت الأساسي — تحويل تلقائي" if st.session_state["lang"] == "ar" else "🔄 Source Sheet — Auto Convert"],
    horizontal=True, key="upload_mode"
)

# ── Source Converter Mode ──
if "تحويل" in _upload_mode or "Convert" in _upload_mode:
    st.markdown(f"""
    <div class="info-box">
    🔄 ارفع الشيت الأساسي (204 عمود) وسيتحول تلقائياً إلى هيكل شيت البرنامج
    </div>
    """, unsafe_allow_html=True)

    _src_up = st.file_uploader(
        "ارفع الشيت الأساسي" if st.session_state["lang"] == "ar" else "Upload Source Sheet",
        type=["xlsx"], key="src_up_main"
    )

    if _src_up:
        try:
            validate_uploaded_file(_src_up)
            _src_b = _src_up.read()

            if st.button("🔄 تحويل وتحليل" if st.session_state["lang"] == "ar" else "🔄 Convert & Analyze",
                         key="convert_main", use_container_width=False):
                with st.spinner("⚙️ جاري التحويل..."):
                    import warnings as _ww
                    _ww.filterwarnings("ignore")

                    _dfs = pd.read_excel(io.BytesIO(_src_b), header=None, skiprows=4)
                    _dfs.columns = range(len(_dfs.columns))

                    # ── وقّف عند آخر صف فيه كود مجمع أو تاريخ زيارة ──
                    _last_code = _dfs[15].last_valid_index()  # الكود المجمع
                    _last_date = _dfs[2].last_valid_index()   # تاريخ الزيارة
                    _last_row  = max(x for x in [_last_code, _last_date] if x is not None)
                    _dfs = _dfs.iloc[:_last_row + 1].copy()

                    def _c(i): return _dfs[i-1] if (i-1) in _dfs.columns else pd.Series([None]*len(_dfs))
                    def _yq(yn, qty):
                        r = []
                        for a, b in zip(yn, qty):
                            if str(a).strip() == 'نعم':
                                try:    r.append(int(float(b)) if pd.notna(b) else None)
                                except: r.append(None)
                            else: r.append(None)
                        return r
                    def _ys(s): return s.apply(lambda x: 'نعم' if str(x).strip()=='نعم' else ('لا' if pd.notna(x) and str(x).strip() not in ['nan',''] else None))

                    _cv = pd.DataFrame()
                    _cv['مسلسل']                                         = _c(1)
                    _cv['الكود المجمع']                                  = _c(16)
                    _cv['اسم الجمعية']                                   = _c(2)
                    _cv['تاريخ الزيارة']                                 = format_idus_source_date(_c(3), swap_excel_dates=False)
                    _cv['اسم مندوب الوصول ثلاثي\n(الباحث الميداني) ']  = _c(4)
                    _cv['محافظة الوصول للمستفيد']                       = _c(5)
                    _cv['منطقة الوصول للمستفيد ']                       = _c(6)
                    _cv['مكان المقابلة']                                 = _c(17)
                    _cv['السن']                                          = _c(18)
                    _cv['النوع']                                         = _c(19)
                    _cv['واقيات']                                        = _yq(_c(27), _c(28))
                    _cv['مزلقات']                                        = _yq(_c(31), _c(32))
                    _cv[' زهري']                                         = _ys(_c(34))
                    _cv['دعم نفسي']                                      = _ys(_c(36))
                    _cv['سرنجات']                                        = _yq(_c(39), _c(40))
                    _cv['ميثادون']                                       = _ys(_c(42))
                    _cv['نتيجة التحليل ']                                = _c(48)
                    _cv['نتيجة التحليل التاكيدي']                       = _c(50)
                    _cv['الاحالة علي العلاج']                           = _c(51)
                    # زيارة متابعة 1 (أعمدة 57-85)
                    _cv['زيارة متابعة 1']        = format_idus_source_date(_c(57))
                    _cv['واقيات متابعة 1']       = _yq(_c(62), _c(63))
                    _cv['مزلقات متابعة 1']       = _yq(_c(66), _c(67))
                    _cv['زهري متابعة 1']         = _ys(_c(69))
                    _cv['دعم نفسي متابعة 1']     = _ys(_c(71))
                    _cv['سرنجات متابعة 1']       = _yq(_c(74), _c(75))
                    _cv['نتيجة تحليل متابعة 1']  = _c(78)
                    _cv['ميثادون متابعة 1']      = _ys(_c(80))
                    # زيارة متابعة 2 (أعمدة 86-114)
                    _cv['زيارة متابعة 2']        = format_idus_source_date(_c(86))
                    _cv['واقيات متابعة 2']       = _yq(_c(91), _c(92))
                    _cv['مزلقات متابعة 2']       = _yq(_c(95), _c(96))
                    _cv['زهري متابعة 2']         = _ys(_c(98))
                    _cv['دعم نفسي متابعة 2']     = _ys(_c(100))
                    _cv['سرنجات متابعة 2']       = _yq(_c(103), _c(104))
                    _cv['نتيجة تحليل متابعة 2']  = _c(107)
                    # زيارة متابعة 3 (أعمدة 115-143)
                    _cv['زيارة متابعة 3']        = format_idus_source_date(_c(115))
                    _cv['واقيات متابعة 3']       = _yq(_c(120), _c(121))
                    _cv['مزلقات متابعة 3']       = _yq(_c(124), _c(125))
                    _cv['نتيجة تحليل متابعة 3']  = _c(136)
                    # زيارة متابعة 4 (أعمدة 144-172)
                    _cv['زيارة متابعة 4']        = format_idus_source_date(_c(144))
                    _cv['نتيجة تحليل متابعة 4']  = _c(165)
                    # زيارة متابعة 5 (أعمدة 173-201)
                    _cv['زيارة متابعة 5']        = format_idus_source_date(_c(173))
                    _cv['نتيجة تحليل متابعة 5']  = _c(194)
                    _cv = _cv.dropna(how='all').reset_index(drop=True)
                    _cv.columns = [str(c).strip() for c in _cv.columns]
                    _cv, _clean_report = clean_dataframe(_cv)

                    st.session_state["hdf"] = _cv
                    st.session_state["data_clean_report"] = _clean_report
                    st.session_state["standard_df"] = generate_standard_df(_cv)
                    st.session_state["mapper_approved"] = False
                    st.session_state["dq_issues"] = None
                    st.session_state["hiv_file_name"] = f"[محوّل] {_src_up.name}"

                    # Download converted file
                    _buf_cv = io.BytesIO()
                    with pd.ExcelWriter(_buf_cv, engine='openpyxl') as _wr:
                        _cv.to_excel(_wr, index=False, sheet_name='IDUs Database')
                    _buf_cv.seek(0)

                st.success(f"✅ تم التحويل! {len(_cv):,} سجل جاهز للتحليل")
                st.download_button(
                    "⬇️ تحميل الشيت المحوّل",
                    data=_buf_cv.getvalue(),
                    file_name=f"IDUs_Converted_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_cv_main"
                )
                st.rerun()

        except Exception as _sce:
            st.error(f"❌ {_sce}")

    if st.session_state["hdf"] is None:
        st.stop()

# ── Direct Upload Mode ──
else:
    hiv_file = st.file_uploader(t("upload_label"), type=["xlsx"], key="hiv_upload")
    if hiv_file is not None:
        try:
            validate_uploaded_file(hiv_file)
            _file_bytes = hiv_file.getvalue()

            if _looks_like_idus_204_source(_file_bytes):
                hdf_new, _clean_report = _convert_idus_204_source_bytes(_file_bytes)
                st.info("🔄 تم اكتشاف الشيت الأساسي 204 عمود وتحويله تلقائيًا قبل التحليل.")
                _display_name = f"[محوّل تلقائي] {hiv_file.name}"
            else:
                hdf_new, _clean_report = smart_read_excel(io.BytesIO(_file_bytes))
                _display_name = hiv_file.name

            st.session_state["hdf"] = hdf_new
            st.session_state["data_clean_report"] = _clean_report
            st.session_state["standard_df"] = generate_standard_df(hdf_new)
            st.session_state["mapper_approved"] = False
            st.session_state["dq_issues"] = None
            st.session_state["hiv_file_name"] = _display_name
        except Exception as e:
            st.error(f"{t('file_error')}: {e}")

    if st.session_state["hdf"] is None:
        st.markdown(f"""
        <div class="info-box">
        {t('upload_hint')}<br>
        {t('upload_hint2')}
        </div>
        """, unsafe_allow_html=True)
        st.stop()

hdf = st.session_state["hdf"]

# ── Smart Data Cleaner status ──
_clean_report = st.session_state.get("data_clean_report")
if _clean_report is not None:
    _report_df = report_to_dataframe(_clean_report)
    _has_errors = bool((_report_df["النوع"] == "error").any()) if not _report_df.empty else False
    _has_warnings = bool((_report_df["النوع"] == "warning").any()) if not _report_df.empty else False
    _box_icon = "❌" if _has_errors else ("⚠️" if _has_warnings else "✅")
    _box_title = "Smart Data Cleaner — فحص وتنظيف الملف"
    with st.expander(f"{_box_icon} {_box_title}", expanded=_has_errors or _has_warnings):
        st.dataframe(_report_df, use_container_width=True, height=260)
        _mapping_df = mapping_report_to_dataframe(_clean_report)
        if not _mapping_df.empty:
            _summary = get_mapping_summary(_clean_report)
            _accepted_count = int(_summary.get("accepted", 0))
            _verify_count = int(_summary.get("verify", 0))
            _suspicious_count = int(_summary.get("suspicious", 0))
            _unknown_count = int(_summary.get("unknown", 0))
            _review_count = _verify_count + _suspicious_count
            _overall_conf = int(_summary.get("overall_confidence", 0))
            _conf_color = "#6bff8e" if _overall_conf >= 95 else ("#ffb74d" if _overall_conf >= 85 else ("#ff9800" if _overall_conf >= 70 else "#ff6b6b"))

            st.markdown("**Auto Data Mapper — لوحة الثقة قبل الاعتماد**")
            st.markdown(f"""
            <div style="display:grid;grid-template-columns:repeat(5,1fr);gap:12px;margin:0.8rem 0 1rem 0;">
              <div class="metric-card"><div class="label">🟢 Auto Accepted</div><div class="value" style="color:#6bff8e">{_accepted_count}</div><div class="sub">95%+ / Memory</div></div>
              <div class="metric-card"><div class="label">🟡 Verify</div><div class="value" style="color:#ffb74d">{_verify_count}</div><div class="sub">85–94% غالبًا صحيح</div></div>
              <div class="metric-card"><div class="label">🟠 Suspicious</div><div class="value" style="color:#ff9800">{_suspicious_count}</div><div class="sub">70–84% يحتاج انتباه</div></div>
              <div class="metric-card"><div class="label">🔴 Unknown</div><div class="value" style="color:#ff6b6b">{_unknown_count}</div><div class="sub">أقل من 70%</div></div>
              <div class="metric-card"><div class="label">Overall Mapping Confidence</div><div class="value" style="color:{_conf_color}">{_overall_conf}%</div><div class="sub">Based on mapped fields</div></div>
            </div>
            """, unsafe_allow_html=True)

            if getattr(_clean_report, "profile_issues", None):
                st.warning("⚠️ Data Profile Validation اكتشف ربطًا مشكوكًا فيه بناءً على محتوى الأعمدة، وليس الاسم فقط.")

            st.markdown("**Auto Data Mapper — تقرير ربط الأعمدة**")
            st.dataframe(_mapping_df, use_container_width=True, height=260)

            _profile_name = getattr(_clean_report, "organization_profile", "global") or "global"
            st.caption(f"Mapping Profile: {_profile_name} — عند الاعتماد سيتم حفظ الربط في AppData/DataBridge/mapping_memory.json لاستخدامه في المرات القادمة.")

            if st.session_state.get("mapper_approved"):
                st.success("✅ تم اعتماد ربط الأعمدة وحفظه في Mapping Memory. يمكنك الاعتماد على التحليل الحالي.")
                if st.button("↩️ إلغاء الاعتماد وإعادة المراجعة", key="unapprove_mapper"):
                    st.session_state["mapper_approved"] = False
                    st.rerun()
            else:
                st.warning("راجع ربط الأعمدة أولًا، ثم اضغط اعتماد قبل ظهور التحليل والرسومات.")
                approve_label = "✅ اعتماد الربط وتشغيل التحليل"
                if _verify_count or _suspicious_count or _unknown_count:
                    approve_label = f"✅ اعتماد الربط رغم وجود Verify ({_verify_count}) / Suspicious ({_suspicious_count}) / Unknown ({_unknown_count})"
                if st.button(approve_label, key="approve_mapper", use_container_width=True):
                    _saved_mappings = save_mapping_memory_from_report(_clean_report, include_review=True)
                    st.session_state["mapper_approved"] = True
                    st.session_state["saved_mappings_count"] = _saved_mappings
                    st.rerun()
        else:
            st.session_state["mapper_approved"] = True

        if _has_errors:
            st.error("فيه أعمدة أساسية ناقصة. راجع التقرير قبل الاعتماد على الأرقام.")
        elif _has_warnings:
            st.warning("تم تنظيف الملف، لكن فيه تحذيرات قد تؤثر على فلاتر الشهر أو بعض الرسومات.")
        else:
            st.success("تم تنظيف الملف وتجهيزه للتحليل بنجاح.")

if _clean_report is not None and not st.session_state.get("mapper_approved", False):
    st.info("⏸️ التحليل متوقف مؤقتًا لحين اعتماد ربط الأعمدة من تقرير Auto Data Mapper بالأعلى.")
    st.stop()

_month_filter_kind = "all"
_selected_month = None
_from_month = None
_to_month = None

# ── Month Filter ──
date_col_main = next((c for c in hdf.columns if 'تاريخ الزيارة' in c and 'متابعة' not in c), None)
if date_col_main:
    hdf[date_col_main] = pd.to_datetime(hdf[date_col_main], errors='coerce')
    _available_months = sorted(hdf[date_col_main].dropna().dt.to_period('M').unique().astype(str).tolist())

    if _available_months:
        st.markdown("---")
        _mf1, _mf2, _mf3 = st.columns([1, 1, 2])
        with _mf1:
            _month_mode = st.radio(
                "🗓️ فلتر الشهر:" if st.session_state["lang"] == "ar" else "🗓️ Month Filter:",
                ["شهر واحد" if st.session_state["lang"] == "ar" else "Single Month",
                 "نطاق" if st.session_state["lang"] == "ar" else "Range",
                 "الكل" if st.session_state["lang"] == "ar" else "All"],
                horizontal=False, key="month_mode"
            )
        with _mf2:
            if ("شهر" in _month_mode or "Single" in _month_mode):
                _sel_month = st.selectbox(
                    "اختر الشهر:" if st.session_state["lang"] == "ar" else "Select Month:",
                    _available_months, key="sel_single_month"
                )
                _month_filter_kind = "single"
                _selected_month = _sel_month
                hdf = hdf[hdf[date_col_main].dt.to_period('M').astype(str) == _sel_month].copy()
            elif ("نطاق" in _month_mode or "Range" in _month_mode):
                _from_m = st.selectbox("من:" if st.session_state["lang"] == "ar" else "From:", _available_months, key="from_month")
                _to_m   = st.selectbox("إلى:" if st.session_state["lang"] == "ar" else "To:", _available_months,
                                        index=len(_available_months)-1, key="to_month")
                _month_filter_kind = "range"
                _from_month = _from_m
                _to_month = _to_m
                hdf = hdf[
                    (hdf[date_col_main].dt.to_period('M').astype(str) >= _from_m) &
                    (hdf[date_col_main].dt.to_period('M').astype(str) <= _to_m)
                ].copy()
        with _mf3:
            st.markdown(f'<div class="info-box" style="margin-top:1.5rem;">📊 السجلات المعروضة: <b style="color:#7c6aff">{len(hdf):,}</b> من إجمالي <b>{len(st.session_state["hdf"]):,}</b></div>', unsafe_allow_html=True)
        st.markdown("---")

base_visits_total = len(hdf)
total = base_visits_total

# ── KPI Calculations ──
# أعمدة تواريخ زيارات المتابعة 1-5 فقط (تاريخ الزيارة)
fu_date_cols = [c for c in hdf.columns if 'زيارة متابعة' in c and ('تاريخ' in c or c.strip() in ['زيارة متابعة 1','زيارة متابعة 2','زيارة متابعة 3','زيارة متابعة 4','زيارة متابعة 5'])]
# fallback: أي عمود اسمه "زيارة متابعة X"
if not fu_date_cols:
    fu_date_cols = [c for c in hdf.columns if 'زيارة متابعة' in c and 'نتيجة' not in c and 'واقيات' not in c and 'مزلقات' not in c and 'زهري' not in c and 'دعم' not in c and 'سرنجات' not in c and 'ميثادون' not in c]
fu_col = fu_date_cols

test_col = [c for c in hdf.columns if 'نتيجة التحليل' in c and 'متابع' not in c and 'تاكيدي' not in c]
ref_col  = [c for c in hdf.columns if 'الاحالة' in c]

def _followup_date_mask(df: pd.DataFrame, cols: List[str]) -> pd.DataFrame:
    """Return a boolean mask for follow-up visit dates respecting the selected month filter.
    Important: follow-ups are filtered by their own visit dates, not by the original visit date.
    """
    masks = []
    for col in cols:
        dt = pd.to_datetime(df[col], errors='coerce')
        if _month_filter_kind == "single" and _selected_month:
            masks.append(dt.dt.to_period('M').astype(str).eq(_selected_month))
        elif _month_filter_kind == "range" and _from_month and _to_month:
            period_s = dt.dt.to_period('M').astype(str)
            masks.append(period_s.ge(_from_month) & period_s.le(_to_month))
        else:
            masks.append(dt.notna())
    if not masks:
        return pd.DataFrame(index=df.index)
    return pd.concat(masks, axis=1)

# إجمالي زيارات المتابعة = عدد تواريخ المتابعة المعبأة
# عند اختيار شهر/نطاق: العد يتم على تاريخ المتابعة نفسه من الشيت الكامل، وليس على تاريخ الزيارة الأساسي.
if fu_col:
    _fu_source_df = st.session_state["hdf"] if _month_filter_kind in ["single", "range"] else hdf
    _fu_df = _followup_date_mask(_fu_source_df, fu_col)
    total_fu_visits  = int(_fu_df.sum().sum())          # إجمالي زيارات المتابعة داخل الشهر/النطاق المختار
    has_followup     = int(_fu_df.any(axis=1).sum())    # مستفيدين عندهم متابعة داخل الشهر/النطاق المختار
else:
    total_fu_visits = has_followup = 0

# في حالة فلتر شهر/نطاق، إجمالي الشهر المطلوب = الزيارات الأساسية داخل الشهر + زيارات المتابعة داخل الشهر
# مثال مايو: 213 زيارة أساسية + 12 متابعة = 225 إجمالي استلام/زيارة.
if _month_filter_kind in ["single", "range"]:
    total = base_visits_total + total_fu_visits
else:
    total = base_visits_total

no_followup = max(base_visits_total - has_followup, 0)
positive     = 0
refused_test = 0
referrals    = 0
if test_col:
    positive     = (hdf[test_col[0]].astype(str).str.strip() == 'ايجابي').sum()
    refused_test = hdf[test_col[0]].isna().sum()
if ref_col:
    referrals = hdf[ref_col[0]].notna().sum()

# ── Visit-level chart dataset (base visits + follow-up visits) ──
def _norm_ar(text: Any) -> str:
    return str(text).replace('ة', 'ه').strip()

def _find_service_col(cols: List[str], keyword: str, followup_no: Optional[int] = None) -> Optional[str]:
    """Find a service/result column, safely handling متابعة/متابعه spelling."""
    norm_kw = _norm_ar(keyword)
    for c in cols:
        nc = _norm_ar(c)
        if norm_kw not in nc:
            continue
        if followup_no is None:
            if 'متابع' not in nc:
                return c
        else:
            if 'متابع' in nc and str(followup_no) in nc:
                return c
    return None

def _service_delivered(value: Any, numeric_service: bool = False) -> bool:
    if pd.isna(value):
        return False
    if numeric_service:
        try:
            return float(value) > 0
        except Exception:
            return False
    return str(value).strip() == 'نعم'

def _date_in_selected_period(dt_series: pd.Series) -> pd.Series:
    dt = pd.to_datetime(dt_series, errors='coerce')
    if _month_filter_kind == "single" and _selected_month:
        return dt.dt.to_period('M').astype(str).eq(_selected_month)
    if _month_filter_kind == "range" and _from_month and _to_month:
        ps = dt.dt.to_period('M').astype(str)
        return ps.ge(_from_month) & ps.le(_to_month)
    return dt.notna()

def build_visit_chart_df(source_df: pd.DataFrame) -> pd.DataFrame:
    """Create one row per visit for charts.
    This prevents charts from ignoring follow-ups when a month filter is active.
    """
    cols = list(source_df.columns)
    age_col0  = next((c for c in cols if 'السن' in c), None)
    gender_col0 = next((c for c in cols if 'النوع' in c or 'الجنس' in c), None)
    area_col0 = next((c for c in cols if 'منطقة' in c or 'منطقه' in c), None)
    gov_col0  = next((c for c in cols if 'محافظة' in c), None)
    main_date = next((c for c in cols if 'تاريخ الزيارة' in c and 'متابعة' not in c and 'متابعه' not in c), None)

    services = [
        ('واقيات', True),
        ('مزلقات', True),
        ('سرنجات', True),
        ('زهري', False),
        ('دعم نفسي', False),
        ('ميثادون', False),
    ]

    rows = []
    def _append_rows(df_part: pd.DataFrame, visit_type: str, date_col: str, result_col: Optional[str], service_cols: Dict[str, Tuple[Optional[str], bool]]):
        for idx, row in df_part.iterrows():
            rec = {
                'نوع الزيارة': visit_type,
                'تاريخ الزيارة': pd.to_datetime(row.get(date_col), errors='coerce'),
                'الشهر': pd.to_datetime(row.get(date_col), errors='coerce').to_period('M').strftime('%Y-%m') if pd.notna(pd.to_datetime(row.get(date_col), errors='coerce')) else None,
                'الفئة العمرية': row.get(age_col0) if age_col0 else None,
                'النوع': row.get(gender_col0) if gender_col0 else None,
                'المنطقة': row.get(area_col0) if area_col0 else None,
                'المحافظة': row.get(gov_col0) if gov_col0 else None,
                'نتيجة التحليل': row.get(result_col) if result_col else None,
            }
            for label, (svc_col, is_num) in service_cols.items():
                rec[label] = _service_delivered(row.get(svc_col), is_num) if svc_col else False
            rows.append(rec)

    # Base visits in the selected period
    if main_date:
        base_mask = _date_in_selected_period(source_df[main_date])
        base_service_cols = {label: (_find_service_col(cols, label, None), is_num) for label, is_num in services}
        base_result_col = _find_service_col(cols, 'نتيجة التحليل', None)
        _append_rows(source_df[base_mask].copy(), 'أساسية', main_date, base_result_col, base_service_cols)

    # Follow-up visits in the selected period, using each follow-up's own date
    for n in range(1, 6):
        fu_date = next((c for c in cols if _norm_ar('زيارة متابعة') in _norm_ar(c) and str(n) in _norm_ar(c)
                        and not any(_norm_ar(x) in _norm_ar(c) for x in ['واقيات','مزلقات','زهري','دعم','سرنجات','ميثادون','نتيجة'])), None)
        if not fu_date:
            continue
        fu_mask = _date_in_selected_period(source_df[fu_date])
        if not fu_mask.any():
            continue
        fu_service_cols = {label: (_find_service_col(cols, label, n), is_num) for label, is_num in services}
        fu_result_col = _find_service_col(cols, 'نتيجة تحليل', n)
        _append_rows(source_df[fu_mask].copy(), f'متابعة {n}', fu_date, fu_result_col, fu_service_cols)

    return pd.DataFrame(rows)

# Charts use the full source + period masks, so follow-ups are counted even when
# the original/base visit happened before the selected month.
visit_chart_df = build_visit_chart_df(st.session_state["hdf"])


# ── Monthly summary + export helpers ──
def _num_sum(series: pd.Series) -> int:
    return int(pd.to_numeric(series, errors='coerce').fillna(0).sum())

def _clean_text_series(series: pd.Series) -> pd.Series:
    return series.fillna('').astype(str).str.strip()

def _find_any_col(df: pd.DataFrame, include_words: List[str], exclude_words: List[str] = []) -> Optional[str]:
    for c in df.columns:
        nc = _norm_ar(c)
        if all(_norm_ar(w) in nc for w in include_words) and not any(_norm_ar(w) in nc for w in exclude_words):
            return c
    return None

def _gender_is_male(value: Any) -> bool:
    s = str(value).strip().replace(' ', '')
    return s in ['ذكر', 'male', 'Male', 'M']

def _gender_is_female(value: Any) -> bool:
    s = str(value).strip().replace(' ', '')
    return s in ['أنثى', 'انثى', 'أنثي', 'انثي', 'female', 'Female', 'F']

def _followup_date_columns_for_df(df: pd.DataFrame) -> List[str]:
    cols = []
    for c in df.columns:
        nc = _norm_ar(c)
        if _norm_ar('زيارة متابعة') in nc and not any(_norm_ar(x) in nc for x in ['واقيات','مزلقات','زهري','دعم','سرنجات','ميثادون','نتيجة']):
            cols.append(c)
    return cols

def get_period_export_df(base_filtered_df: pd.DataFrame, full_df: pd.DataFrame) -> pd.DataFrame:
    """Export rows relevant to the selected period.
    With a month/range filter, include rows with either a base visit in the period
    OR any follow-up visit date in the period. This prevents follow-up columns from
    exporting empty when the base visit was in a previous month.
    """
    if _month_filter_kind not in ['single', 'range']:
        return base_filtered_df.copy()

    mask = pd.Series(False, index=full_df.index)
    main_date = next((c for c in full_df.columns if 'تاريخ الزيارة' in c and 'متابعة' not in c and 'متابعه' not in c), None)
    if main_date:
        mask |= _date_in_selected_period(full_df[main_date])
    for c in _followup_date_columns_for_df(full_df):
        mask |= _date_in_selected_period(full_df[c])
    out = full_df[mask].copy()
    return out.reset_index(drop=True)

def build_monthly_summary_tables(base_df: pd.DataFrame, full_df: pd.DataFrame, visit_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """Build the detailed report and the compact 'جدول مثل الصورة' from the current filters."""
    gender_col = _find_any_col(base_df, ['النوع'])
    male_count = int(base_df[gender_col].apply(_gender_is_male).sum()) if gender_col else 0
    female_count = int(base_df[gender_col].apply(_gender_is_female).sum()) if gender_col else 0

    # Base quantities from base visits in the selected period
    base_syr_col = _find_any_col(base_df, ['سرنجات'], ['متابع'])
    base_cond_col = _find_any_col(base_df, ['واقيات'], ['متابع'])
    base_lube_col = _find_any_col(base_df, ['مزلقات'], ['متابع'])
    base_syr = _num_sum(base_df[base_syr_col]) if base_syr_col else 0
    base_cond = _num_sum(base_df[base_cond_col]) if base_cond_col else 0
    base_lube = _num_sum(base_df[base_lube_col]) if base_lube_col else 0

    # Follow-up quantities by each follow-up's own date, from the full source sheet
    fu_syr = fu_cond = fu_lube = 0
    for n in range(1, 6):
        fu_date = next((c for c in full_df.columns if _norm_ar('زيارة متابعة') in _norm_ar(c) and str(n) in _norm_ar(c)
                        and not any(_norm_ar(x) in _norm_ar(c) for x in ['واقيات','مزلقات','زهري','دعم','سرنجات','ميثادون','نتيجة'])), None)
        if not fu_date:
            continue
        mask = _date_in_selected_period(full_df[fu_date])
        syr_col = _find_service_col(list(full_df.columns), 'سرنجات', n)
        cond_col = _find_service_col(list(full_df.columns), 'واقيات', n)
        lube_col = _find_service_col(list(full_df.columns), 'مزلقات', n)
        fu_syr += _num_sum(full_df.loc[mask, syr_col]) if syr_col else 0
        fu_cond += _num_sum(full_df.loc[mask, cond_col]) if cond_col else 0
        fu_lube += _num_sum(full_df.loc[mask, lube_col]) if lube_col else 0

    # Male/female recipients and totals by commodity, using the visit-level table for recipient counts
    def _recipients_by_gender(tool: str, male: bool) -> int:
        """Count recipients by gender from the visit-level table when available."""
        if visit_df.empty or tool not in visit_df.columns or 'النوع' not in visit_df.columns:
            return 0
        delivered = visit_df[tool].fillna(False).astype(bool)
        gender_mask = visit_df['النوع'].apply(_gender_is_male if male else _gender_is_female)
        return int((delivered & gender_mask).sum())

    # Better gender split for quantities/recipients directly from source rows
    def _base_recipient_count(tool_col: Optional[str], want_male: bool) -> int:
        if not tool_col or not gender_col:
            return 0
        vals = pd.to_numeric(base_df[tool_col], errors='coerce').fillna(0) > 0
        genders = base_df[gender_col].apply(_gender_is_male if want_male else _gender_is_female)
        return int((vals & genders).sum())

    def _base_qty_by_gender(tool_col: Optional[str], want_male: bool) -> int:
        if not tool_col or not gender_col:
            return 0
        genders = base_df[gender_col].apply(_gender_is_male if want_male else _gender_is_female)
        return _num_sum(base_df.loc[genders, tool_col])

    def _fu_recipient_count(tool: str, want_male: bool) -> int:
        gcol_full = _find_any_col(full_df, ['النوع'])
        if not gcol_full:
            return 0
        count = 0
        for n in range(1, 6):
            fu_date = next((c for c in full_df.columns if _norm_ar('زيارة متابعة') in _norm_ar(c) and str(n) in _norm_ar(c)
                            and not any(_norm_ar(x) in _norm_ar(c) for x in ['واقيات','مزلقات','زهري','دعم','سرنجات','ميثادون','نتيجة'])), None)
            tcol = _find_service_col(list(full_df.columns), tool, n)
            if not fu_date or not tcol:
                continue
            mask = _date_in_selected_period(full_df[fu_date])
            got = pd.to_numeric(full_df[tcol], errors='coerce').fillna(0) > 0
            genders = full_df[gcol_full].apply(_gender_is_male if want_male else _gender_is_female)
            count += int((mask & got & genders).sum())
        return count

    def _fu_qty_by_gender(tool: str, want_male: bool) -> int:
        gcol_full = _find_any_col(full_df, ['النوع'])
        if not gcol_full:
            return 0
        total_qty = 0
        for n in range(1, 6):
            fu_date = next((c for c in full_df.columns if _norm_ar('زيارة متابعة') in _norm_ar(c) and str(n) in _norm_ar(c)
                            and not any(_norm_ar(x) in _norm_ar(c) for x in ['واقيات','مزلقات','زهري','دعم','سرنجات','ميثادون','نتيجة'])), None)
            tcol = _find_service_col(list(full_df.columns), tool, n)
            if not fu_date or not tcol:
                continue
            mask = _date_in_selected_period(full_df[fu_date])
            genders = full_df[gcol_full].apply(_gender_is_male if want_male else _gender_is_female)
            total_qty += _num_sum(full_df.loc[mask & genders, tcol])
        return total_qty

    male_syr_n = _base_recipient_count(base_syr_col, True) + _fu_recipient_count('سرنجات', True)
    male_cond_n = _base_recipient_count(base_cond_col, True) + _fu_recipient_count('واقيات', True)
    male_lube_n = _base_recipient_count(base_lube_col, True) + _fu_recipient_count('مزلقات', True)
    female_syr_n = _base_recipient_count(base_syr_col, False) + _fu_recipient_count('سرنجات', False)
    female_cond_n = _base_recipient_count(base_cond_col, False) + _fu_recipient_count('واقيات', False)
    female_lube_n = _base_recipient_count(base_lube_col, False) + _fu_recipient_count('مزلقات', False)
    male_syr_qty = _base_qty_by_gender(base_syr_col, True) + _fu_qty_by_gender('سرنجات', True)
    male_cond_qty = _base_qty_by_gender(base_cond_col, True) + _fu_qty_by_gender('واقيات', True)
    male_lube_qty = _base_qty_by_gender(base_lube_col, True) + _fu_qty_by_gender('مزلقات', True)
    female_syr_qty = _base_qty_by_gender(base_syr_col, False) + _fu_qty_by_gender('سرنجات', False)
    female_cond_qty = _base_qty_by_gender(base_cond_col, False) + _fu_qty_by_gender('واقيات', False)
    female_lube_qty = _base_qty_by_gender(base_lube_col, False) + _fu_qty_by_gender('مزلقات', False)

    # Tests and service counts from all visits in the period
    if not visit_df.empty and 'نتيجة التحليل' in visit_df.columns:
        res = _clean_text_series(visit_df['نتيجة التحليل'])
        res = res[~res.isin(['', 'nan', 'None'])]
        total_tests = int(len(res))
        negative_tests = int(res.str.replace('إ', 'ا').eq('سلبي').sum())
        positive_tests = int(res.str.replace('إ', 'ا').eq('ايجابي').sum())
    else:
        total_tests = negative_tests = positive_tests = 0

    syphilis_count = int(visit_df['زهري'].fillna(False).astype(bool).sum()) if not visit_df.empty and 'زهري' in visit_df.columns else 0
    psycho_count = int(visit_df['دعم نفسي'].fillna(False).astype(bool).sum()) if not visit_df.empty and 'دعم نفسي' in visit_df.columns else 0
    ost_count = int(visit_df['ميثادون'].fillna(False).astype(bool).sum()) if not visit_df.empty and 'ميثادون' in visit_df.columns else 0

    period_label = _selected_month or ((_from_month + ' إلى ' + _to_month) if _from_month and _to_month else 'كل الفترة')
    detailed_rows = [
        [f'تقرير {period_label} - الأساسي + المتابعة', None, None, None],
        [None, None, None, None],
        ['1) عدد المستفيدين بدون متابعة', None, None, None],
        ['البيان', 'العدد', None, None],
        ['ذكر', male_count, None, None],
        ['أنثى', female_count, None, None],
        ['الإجمالي', base_visits_total, None, None],
        ['إجمالي عدد المتابعة', total_fu_visits, None, None],
        [None, None, None, None],
        [None, None, None, None],
        ['2) إجمالي توزيع بدون متابعة', None, None, None],
        ['الصنف', 'الإجمالي', None, None],
        ['سرنجات', base_syr, None, None],
        ['واقي', base_cond, None, None],
        ['مزلق', base_lube, None, None],
        [None, None, None, None],
        [None, None, None, None],
        ['3) إجمالي توزيع متابعة', None, None, None],
        ['الصنف', 'الإجمالي', None, None],
        ['سرنجات', fu_syr, None, None],
        ['واقي', fu_cond, None, None],
        ['مزلق', fu_lube, None, None],
        [None, None, None, None],
        [None, None, None, None],
        ['4) كم مستفيد ذكر استلم الأساسي + المتابعة', None, None, None],
        ['الصنف', 'عدد الذكور', 'إجمالي الاستلامات', 'ملخص'],
        ['سرنجات', male_syr_n, male_syr_qty, f'{male_syr_n} ذكر استلموا {male_syr_qty:,} سرنجة'],
        ['واقي', male_cond_n, male_cond_qty, f'{male_cond_n} ذكر استلموا {male_cond_qty:,} واقي'],
        ['مزلق', male_lube_n, male_lube_qty, f'{male_lube_n} ذكر استلموا {male_lube_qty:,} مزلق'],
        [None, None, None, None],
        [None, None, None, None],
        ['5) كم مستفيد أنثى استلم الأساسي + المتابعة', None, None, None],
        ['الصنف', 'عدد الإناث', 'إجمالي الاستلامات', 'ملخص'],
        ['سرنجات', female_syr_n, female_syr_qty, f'{female_syr_n} أنثى استلموا {female_syr_qty:,} سرنجة'],
        ['واقي', female_cond_n, female_cond_qty, f'{female_cond_n} أنثى استلموا {female_cond_qty:,} واقي'],
        ['مزلق', female_lube_n, female_lube_qty, f'{female_lube_n} أنثى استلموا {female_lube_qty:,} مزلق'],
        [None, None, None, None],
        [None, None, None, None],
        ['6) عدد التحاليل الأساسي + المتابعة', None, None, None],
        ['البيان', 'العدد', None, None],
        ['إجمالي التحاليل', total_tests, None, None],
        ['سلبي', negative_tests, None, None],
        ['إيجابي', positive_tests, None, None],
        ['الإيجابي حسب النوع', 'غير متوفر', None, None],
        [None, None, None, None],
        [None, None, None, None],
        ['7) عدد الخدمات الأساسي + المتابعة', None, None, None],
        ['الخدمة', 'العدد', None, None],
        ['الزهري', syphilis_count, None, None],
        ['الدعم النفسي', psycho_count, None, None],
        ['بدائل الأفيونات / OST', ost_count, None, None],
    ]
    detailed_df = pd.DataFrame(detailed_rows, columns=['البيان', 'العدد', 'إجمالي الاستلامات', 'ملخص'])

    compact_rows = [
        [base_visits_total + total_fu_visits, 'العدد'],
        [base_visits_total, 'الوصول'],
        [male_count, 'رجال'],
        [total_fu_visits, 'متابعه'],
        [base_syr, 'سرنجات'],
        [fu_syr, 'سرنجات متابعه'],
        [base_lube, 'مزلقات'],
        [fu_lube, 'متابعه مزلق'],
        [base_cond, 'واقيات'],
        [fu_cond, 'متابعة واق'],
        [total_tests, 'اجمالي التحاليل'],
        [total_fu_visits, 'متابعه'],
        [negative_tests, 'سلبي'],
        [positive_tests, 'ايجابي'],
        ['غير متوفر', 'ايجابي تأكيدي'],
        [psycho_count, 'دعم نفسي'],
        [syphilis_count, 'زهري'],
    ]
    compact_df = pd.DataFrame(compact_rows, columns=['العدد', 'البيان'])
    return detailed_df, compact_df

summary_report_df, picture_table_df = build_monthly_summary_tables(hdf, st.session_state['hdf'], visit_chart_df)
export_hdf = get_period_export_df(hdf, st.session_state['hdf'])

# ── Quality Score Banner (auto-run on file load) ──

qs = None
if st.session_state["dq_issues"] is not None:
    qs = compute_quality_score(st.session_state["dq_issues"], total)
    # Gauge bar
    bar_pct = qs['score']
    st.markdown(f"""
    <div style="background:#0d0d1a;border:1px solid #1e1e2e;border-radius:14px;
                padding:1.2rem 1.8rem;margin-bottom:1.5rem;
                display:flex;align-items:center;gap:2rem;flex-wrap:wrap;">
      <div>
        <div style="font-size:0.72rem;color:#555;text-transform:uppercase;letter-spacing:0.1em;">
          Data Quality Score
        </div>
        <div style="font-size:2.8rem;font-weight:800;color:{qs['color']};
                    font-family:'JetBrains Mono',monospace;line-height:1.1;">
          {qs['score']}%
        </div>
        <div style="font-size:0.85rem;color:{qs['color']};font-weight:600;">{qs['status']}</div>
      </div>
      <div style="flex:1;min-width:200px;">
        <div style="background:#1a1a2e;border-radius:8px;height:12px;overflow:hidden;margin-bottom:0.8rem;">
          <div style="background:linear-gradient(90deg,{qs['color']}88,{qs['color']});
                      width:{bar_pct}%;height:100%;border-radius:8px;
                      transition:width 0.5s ease;"></div>
        </div>
        <div style="display:flex;gap:1.5rem;flex-wrap:wrap;">
          <span style="font-size:0.78rem;">🔴 Critical <b style="color:#ff6b6b">{qs['critical']}</b></span>
          <span style="font-size:0.78rem;">🟠 Major <b style="color:#ff9800">{qs['major']}</b></span>
          <span style="font-size:0.78rem;">🟡 Minor <b style="color:#ffb74d">{qs['minor']}</b></span>
          <span style="font-size:0.78rem;">ℹ️ Info <b style="color:#7c6aff">{qs['info']}</b></span>
          <span style="font-size:0.78rem;">🟢 Valid <b style="color:#6bff8e">{qs['valid']:,}</b></span>
          <span style="font-size:0.78rem;color:#555;">Total <b style="color:#888">{total:,}</b></span>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

# ── KPI Dashboard ──
st.markdown(f"""
<div class="metrics-row">
  <div class="metric-card">
    <div class="label">{'إجمالي الشهر' if _month_filter_kind in ['single','range'] and st.session_state['lang']=='ar' else t('total')}</div>
    <div class="value">{total:,}</div>
    <div class="sub">{f'{base_visits_total:,} أساسي + {total_fu_visits:,} متابعة' if _month_filter_kind in ['single','range'] and st.session_state['lang']=='ar' else t('beneficiary')}</div>
  </div>
  <div class="metric-card">
    <div class="label">{t('has_followup')}</div>
    <div class="value">{has_followup}</div>
    <div class="sub">{round(has_followup/total*100,1) if total>0 else 0}% {t('from_total')}</div>
  </div>
  <div class="metric-card">
    <div class="label">{'إجمالي زيارات المتابعة' if st.session_state['lang']=='ar' else 'Total Follow-up Visits'}</div>
    <div class="value" style="color:#6bb3ff">{total_fu_visits}</div>
    <div class="sub">{'زيارات 1 - 5' if st.session_state['lang']=='ar' else 'Visits 1 - 5'}</div>
  </div>
  <div class="metric-card">
    <div class="label">{t('positive')}</div>
    <div class="value" style="color:#ff6b6b">{positive}</div>
    <div class="sub">{t('positive_case')}</div>
  </div>
  <div class="metric-card">
    <div class="label">{t('refused')}</div>
    <div class="value">{refused_test}</div>
    <div class="sub">{t('beneficiary')}</div>
  </div>
</div>
""", unsafe_allow_html=True)

# ════════════════════════════════════════
#  TABS
# ════════════════════════════════════════
_summary_tab_label = (
    f"📋 ملخص شهر {_selected_month}"
    if _month_filter_kind == "single" and _selected_month
    else "📋 ملخص البيانات"
)
tab_data, tab_repair, tab_gaps, tab_stats, tab_summary, tab_export, tab_quality, tab_indicators, tab_ai, tab_settings, tab_converter = st.tabs([
    t("tab_data"),
    "🛠 مركز إصلاح البيانات",
    t("tab_gaps"),
    t("tab_stats"),
    _summary_tab_label,
    t("tab_export"),
    t("tab_quality"),
    "📈 Indicators",
    "🤖 AI Assistant",
    "⚙️ Settings",
    "🔄 Source Converter",
])

# ──────────────────────────────────────
#  TAB 1 — DATA
# ──────────────────────────────────────
with tab_data:
    show_mode = st.radio(t("show_label"), [
        t("all_data"),
        t("without_followup"),
        t("with_followup"),
        t("positive_cases")
    ], horizontal=True)

    if fu_col:
        if show_mode == t("without_followup"):
            display_hdf = hdf[hdf[fu_col[0]].isna()].copy()
        elif show_mode == t("with_followup"):
            display_hdf = hdf[hdf[fu_col[0]].notna()].copy()
        elif show_mode == t("positive_cases") and test_col:
            display_hdf = hdf[hdf[test_col[0]].astype(str).str.strip() == 'ايجابي'].copy()
        else:
            display_hdf = hdf.copy()
    else:
        display_hdf = hdf.copy()

    st.markdown(f'<div class="info-box">{t("rows_count")}: <b style="color:#7c6aff">{len(display_hdf):,}</b></div>', unsafe_allow_html=True)
    st.dataframe(display_hdf, use_container_width=True, height=420)


# ──────────────────────────────────────
#  TAB 2 — GAP ANALYSIS
# ──────────────────────────────────────
with tab_gaps:
    st.markdown(f'<div class="info-box">{t("gaps_info")}</div>', unsafe_allow_html=True)

    conditional_cols = {
        'نتيجة التحليل التاكيدي':  t("col_reason_confirm"),
        'الاحالة علي العلاج':      t("col_reason_referral"),
        'نتيجة التحليل متابعه':    t("col_reason_test_fu"),
        'واقيات متابعه':           t("col_reason_condoms_fu"),
        'مزلقات متابعه':           t("col_reason_lube_fu"),
        'زهري متابعه':             t("col_reason_syph_fu"),
        'الدعم النفسي متابعه':     t("col_reason_psycho_fu"),
        'سرنجات متابعه':           t("col_reason_syr_fu"),
        'ميثادون متابعه':          t("col_reason_meth_fu"),
    }

    null_reasons = []
    for col_name, reason in conditional_cols.items():
        matched = [c for c in hdf.columns if col_name in c]
        if matched:
            n_null = hdf[matched[0]].isna().sum()
            if n_null > 0:
                null_reasons.append({
                    t("col_name"):   matched[0],
                    t("gap_count"):  n_null,
                    t("gap_pct"):    f"{round(n_null/total*100,1)}%",
                    t("gap_reason"): reason,
                    t("gap_status"): t("justified"),
                })

    other_nulls = []
    justified_cols = [r[t("col_name")] for r in null_reasons]
    for col in hdf.columns:
        n_null = hdf[col].isna().sum()
        if n_null > 0 and col not in justified_cols:
            other_nulls.append({
                t("col_name"):   col,
                t("gap_count"):  n_null,
                t("gap_pct"):    f"{round(n_null/total*100,1)}%",
                t("gap_reason"): t("col_reason_default"),
                t("gap_status"): t("needs_review"),
            })

    all_nulls = null_reasons + other_nulls
    if all_nulls:
        st.subheader(t("gaps_report"))
        st.dataframe(pd.DataFrame(all_nulls), use_container_width=True, height=380)
    else:
        st.success(t("no_gaps"))



# ── Chart label helpers ──
def show_bar_values(fig):
    """Show numeric values on bars so exported chart images keep the numbers visible."""
    fig.update_traces(texttemplate='%{y}', textposition='outside', cliponaxis=False)
    fig.update_layout(uniformtext_minsize=10, uniformtext_mode='show')
    return fig

def show_pie_values(fig):
    """Show label + value + percent on pie charts."""
    fig.update_traces(textinfo='label+value+percent', textposition='auto')
    return fig

def show_line_values(fig):
    """Show numeric values on line chart markers."""
    fig.update_traces(mode='lines+markers+text', texttemplate='%{y}', textposition='top center')
    return fig

# ──────────────────────────────────────
#  TAB 3 — STATISTICS
# ──────────────────────────────────────
with tab_stats:
    ch1, ch2 = st.columns(2)

    with ch1:
        if not visit_chart_df.empty and 'نتيجة التحليل' in visit_chart_df.columns:
            test_counts = visit_chart_df['نتيجة التحليل'].fillna(t("refused_test")).value_counts().reset_index()
            test_counts.columns = ['النتيجة', 'العدد']
            fig = px.pie(test_counts, names='النتيجة', values='العدد',
                        title=t("chart_test"), template='plotly_dark',
                        color_discrete_sequence=['#7c6aff','#ff6b6b','#6bffb8'])
            fig.update_layout(paper_bgcolor='rgba(0,0,0,0)')
            fig = show_pie_values(fig)
            st.plotly_chart(fig, use_container_width=True)

    with ch2:
        # عند فلترة شهر/نطاق: هذا الرسم يجب أن يعرض إجمالي الشهر كـ
        # زيارات أساسية داخل الشهر + زيارات متابعة داخل الشهر، وليس "بدون متابعة".
        if _month_filter_kind in ["single", "range"]:
            fu_df = pd.DataFrame({
                'الحالة': [
                    'زيارات أساسية' if st.session_state["lang"] == 'ar' else 'Base visits',
                    'زيارات متابعة' if st.session_state["lang"] == 'ar' else 'Follow-up visits'
                ],
                'العدد': [base_visits_total, total_fu_visits]
            })
            _fu_title = 'توزيع زيارات الشهر' if st.session_state["lang"] == 'ar' else 'Monthly Visit Distribution'
        else:
            fu_df = pd.DataFrame({
                'الحالة': [t("with_followup_label"), t("without_followup_label")],
                'العدد': [has_followup, no_followup]
            })
            _fu_title = t("chart_followup")
        fig2 = px.bar(fu_df, x='الحالة', y='العدد',
                     title=_fu_title, template='plotly_dark', text='العدد',
                     color_discrete_sequence=['#7c6aff'])
        fig2.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(13,13,26,1)')
        fig2 = show_bar_values(fig2)
        st.plotly_chart(fig2, use_container_width=True)

    ch3, ch4 = st.columns(2)

    with ch3:
        svc_data = {}
        if not visit_chart_df.empty:
            for label in ['زهري', 'دعم نفسي', 'ميثادون', 'واقيات', 'مزلقات', 'سرنجات']:
                if label in visit_chart_df.columns:
                    svc_data[label] = int(visit_chart_df[label].fillna(False).astype(bool).sum())
        if svc_data:
            svc_df = pd.DataFrame({'الخدمة': list(svc_data.keys()), 'العدد': list(svc_data.values())})
            fig3 = px.bar(svc_df, x='الخدمة', y='العدد',
                         title=t("chart_services"), template='plotly_dark',
                         color_discrete_sequence=['#7c6aff'])
            fig3.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(13,13,26,1)')
            fig3 = show_bar_values(fig3)
            st.plotly_chart(fig3, use_container_width=True)

    with ch4:
        age_col = ['الفئة العمرية'] if (not visit_chart_df.empty and 'الفئة العمرية' in visit_chart_df.columns) else []
        if age_col:
            age_counts = visit_chart_df['الفئة العمرية'].dropna().value_counts().reset_index()
            age_counts.columns = ['الفئة العمرية', 'العدد']
            fig4 = px.bar(age_counts, x='الفئة العمرية', y='العدد',
                         title=t("chart_age"), template='plotly_dark',
                         color_discrete_sequence=['#6bffb8'])
            fig4.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(13,13,26,1)')
            fig4 = show_bar_values(fig4)
            st.plotly_chart(fig4, use_container_width=True)

    # Geographic distribution
    area_col = ['المنطقة'] if (not visit_chart_df.empty and 'المنطقة' in visit_chart_df.columns) else []
    if area_col:
        ch5, ch6 = st.columns(2)
        with ch5:
            area_counts = visit_chart_df['المنطقة'].dropna().value_counts().reset_index()
            area_counts.columns = ['المنطقة', 'العدد']
            fig_area = px.bar(area_counts, x='المنطقة', y='العدد',
                             title=t("chart_area"), template='plotly_dark',
                             color_discrete_sequence=['#ff6b9d'])
            fig_area.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(13,13,26,1)')
            fig_area = show_bar_values(fig_area)
            st.plotly_chart(fig_area, use_container_width=True)

        with ch6:
            gov_col = ['المحافظة'] if (not visit_chart_df.empty and 'المحافظة' in visit_chart_df.columns) else []
            if gov_col:
                gov_counts = visit_chart_df['المحافظة'].dropna().value_counts().reset_index()
                gov_counts.columns = ['المحافظة', 'العدد']
                fig_gov = px.pie(gov_counts, names='المحافظة', values='العدد',
                                title=t("chart_gov"), template='plotly_dark',
                                color_discrete_sequence=px.colors.sequential.Purples_r)
                fig_gov.update_layout(paper_bgcolor='rgba(0,0,0,0)')
                fig_gov = show_pie_values(fig_gov)
                st.plotly_chart(fig_gov, use_container_width=True)

    # Age + protective tools
    if age_col:
        st.subheader(t("chart_age_prot"))
        prot_data = []
        for label in ['واقيات', 'مزلقات', 'سرنجات']:
            if label in visit_chart_df.columns and 'الفئة العمرية' in visit_chart_df.columns:
                tmp = visit_chart_df[visit_chart_df[label].fillna(False).astype(bool)]
                tmp = tmp.groupby('الفئة العمرية').size().reset_index(name='العدد')
                tmp['الأداة'] = label
                tmp.columns = ['الفئة العمرية', 'العدد', 'الأداة']
                prot_data.append(tmp)
        if prot_data:
            prot_df = pd.concat(prot_data)
            fig_prot = px.bar(prot_df, x='الفئة العمرية', y='العدد',
                             color='الأداة', barmode='group',
                             title=t("chart_age_prot"), template='plotly_dark',
                             color_discrete_sequence=['#7c6aff','#6bffb8','#ff6b9d'])
            fig_prot.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(13,13,26,1)')
            fig_prot = show_bar_values(fig_prot)
            st.plotly_chart(fig_prot, use_container_width=True)

    # Monthly visits: base + follow-up visits
    if not visit_chart_df.empty and 'الشهر' in visit_chart_df.columns:
        monthly = visit_chart_df.dropna(subset=['الشهر']).groupby('الشهر').size().reset_index(name='عدد الزيارات')
        fig5 = px.line(monthly, x='الشهر', y='عدد الزيارات',
                      title=t("chart_monthly"), template='plotly_dark',
                      color_discrete_sequence=['#7c6aff'], markers=True)
        fig5.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(13,13,26,1)')
        fig5 = show_line_values(fig5)
        st.plotly_chart(fig5, use_container_width=True)


# ──────────────────────────────────────
#  TAB 4 — DATA SUMMARY
# ──────────────────────────────────────
with tab_summary:
    _period_label = _selected_month or ((_from_month + ' إلى ' + _to_month) if _from_month and _to_month else 'كل الفترة')
    st.markdown(f"### {_summary_tab_label}")
    st.caption(f"ملخص الفترة: {_period_label}")
    st.dataframe(picture_table_df, use_container_width=True, hide_index=True, height=500)

    _compact_buf = io.BytesIO()
    with pd.ExcelWriter(_compact_buf, engine='openpyxl') as _writer:
        summary_report_df.to_excel(_writer, index=False, sheet_name='تقرير الفترة')
        picture_table_df.to_excel(_writer, index=False, sheet_name='ملخص البيانات')
    st.download_button(
        "⬇️ تحميل ملخص البيانات Excel",
        data=_compact_buf.getvalue(),
        file_name=f"ملخص_البيانات_{_period_label}.xlsx".replace(" ", "_"),
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key="download_picture_table"
    )


# ──────────────────────────────────────
#  TAB 5 — EXPORT
# ──────────────────────────────────────
with tab_export:
    ex1, ex2 = st.columns(2)

    with ex1:
        csv_hiv = export_hdf.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            t("export_csv"), csv_hiv,
            "IDUs_cleaned.csv", "text/csv",
            use_container_width=True
        )

    with ex2:
        buf_hiv = io.BytesIO()
        with pd.ExcelWriter(buf_hiv, engine='openpyxl') as writer:
            # عند وجود فلتر شهر/نطاق، يتم تصدير الصفوف التي لها زيارة أساسية داخل الفترة
            # أو متابعة داخل الفترة، لذلك لن تظهر أعمدة المتابعة فاضية بسبب الفلتر.
            export_hdf.to_excel(writer, index=False, sheet_name='البيانات')
            if all_nulls:
                pd.DataFrame(all_nulls).to_excel(writer, index=False, sheet_name='تقرير الفراغات')
            summary_report_df.to_excel(writer, index=False, sheet_name='تقرير الفترة')
            picture_table_df.to_excel(writer, index=False, sheet_name='ملخص البيانات')
            if not visit_chart_df.empty:
                visit_chart_df.to_excel(writer, index=False, sheet_name='زيارات الفترة')
        st.download_button(
            t("export_excel"), buf_hiv.getvalue(),
            "IDUs_with_report.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

    st.markdown("---")
    st.markdown(f"### {t('save_excel_title')}")

    excel_db_file = st.file_uploader(t("upload_db"), type=["xlsx"], key="excel_save_db")

    if excel_db_file and st.button(t("save_all_btn"), key="save_all_records"):
        try:
            wb = openpyxl.load_workbook(excel_db_file)
            sheet_name = next(
                (n for n in wb.sheetnames if 'IDU' in n.upper() or 'بيانات' in n or 'data' in n.lower()),
                wb.sheetnames[0]
            )
            ws = wb[sheet_name]

            # Find last used row
            last_row = 5
            for row in ws.iter_rows(min_row=5, values_only=True):
                if any(v is not None for v in row):
                    last_row += 1

            # Write each row from hdf
            added = 0
            for _, row_data in hdf.iterrows():
                col_idx = 1
                for val in row_data.values:
                    ws.cell(row=last_row, column=col_idx, value=val)
                    col_idx += 1
                last_row += 1
                added += 1

            output = BytesIO()
            wb.save(output)
            output.seek(0)

            st.success(f"✅ {t('save_success')} ({added} rows)")
            st.download_button(
                t("download_updated"),
                data=output.getvalue(),
                file_name="IDUs_Database_updated.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_excel_final"
            )
        except Exception as e:
            st.error(f"❌ {e}")


# ──────────────────────────────────────
#  TAB 5 — DATA QUALITY ENGINE (Phase 2)
# ──────────────────────────────────────
with tab_quality:
    st.markdown(f"""
    <div style="margin-bottom:1.5rem;">
      <div style="font-size:1.1rem;font-weight:700;color:#e0e0f0;">{t('dq_title')}</div>
      <div style="font-size:0.78rem;color:#555;margin-top:0.3rem;">{t('dq_subtitle')}</div>
    </div>
    """, unsafe_allow_html=True)

    # Run quality check button
    if st.button(t("dq_run_btn"), key="run_dq"):
        with st.spinner("🔍 جاري الفحص..."):
            st.session_state["dq_issues"] = run_quality_engine(hdf)
            st.session_state["dq_notes"]  = {
                i["row_idx"]: st.session_state["dq_notes"].get(i["row_idx"], "")
                for i in st.session_state["dq_issues"]
            }

    issues = st.session_state.get("dq_issues")

    if issues is None:
        st.markdown(f'<div class="info-box">اضغط زر الفحص لتشغيل محرك جودة البيانات على {total:,} سجل</div>', unsafe_allow_html=True)
    else:
        errors   = [i for i in issues if i["severity"] == "error"]
        warnings = [i for i in issues if i["severity"] == "warning"]
        ok_count = total - len({i["row_idx"] for i in issues})
        ok_count = max(ok_count, 0)

        # ── Summary KPIs ──
        st.markdown(f"""
        <div style="display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin-bottom:1.5rem;">
          <div class="metric-card" style="border-color:#ff6b6b44;">
            <div class="label">{t('dq_errors')}</div>
            <div class="value" style="color:#ff6b6b;">🔴 {len(errors)}</div>
          </div>
          <div class="metric-card" style="border-color:#ffb74d44;">
            <div class="label">{t('dq_warnings')}</div>
            <div class="value" style="color:#ffb74d;">🟡 {len(warnings)}</div>
          </div>
          <div class="metric-card" style="border-color:#6bff8e44;">
            <div class="label">{t('dq_ok')}</div>
            <div class="value" style="color:#6bff8e;">🟢 {ok_count:,}</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        if not issues:
            st.success(t("dq_no_issues"))
        else:
            # ── Filter ──
            filter_choice = st.radio(
                t("dq_filter_label"),
                [t("dq_show_all_issues"), t("dq_show_errors_only"), t("dq_show_warnings_only")],
                horizontal=True, key="dq_filter"
            )

            if filter_choice == t("dq_show_errors_only"):
                display_issues = errors
            elif filter_choice == t("dq_show_warnings_only"):
                display_issues = warnings
            else:
                display_issues = issues

            st.markdown(f'<div class="info-box">{t("dq_records_with_issues")}: <b style="color:#ff6b6b">{len({i["row_idx"] for i in display_issues}):,}</b> &nbsp;|&nbsp; إجمالي المشاكل: <b>{len(display_issues)}</b></div>', unsafe_allow_html=True)

            # ── Issues Table with Notes ──
            st.markdown("---")

            # Build display dataframe
            rows_data = []
            for issue in display_issues:
                sev_badge = "🔴" if issue["severity"] == "error" else "🟡"
                rows_data.append({
                    t("dq_col_row"):      issue["row_idx"] + 1,
                    "مسلسل":              issue["serial"],
                    t("dq_col_field"):    issue["field"],
                    t("dq_col_value"):    issue["value"],
                    t("dq_col_rule"):     t(issue["rule_key"]),
                    t("dq_col_severity"): sev_badge,
                })

            issues_df = pd.DataFrame(rows_data)
            st.dataframe(issues_df, use_container_width=True, height=380)

            # ── Notes Section (for under-15 and other flagged records) ──
            st.markdown("---")
            st.markdown(f"### 📝 {t('dq_col_notes')}")

            under15_issues = [i for i in display_issues if i["rule_key"] == "dq_err_age_u15"]
            if under15_issues:
                st.markdown('<div class="info-box">🔴 السجلات التالية تحتاج ملاحظة ولي الأمر:</div>', unsafe_allow_html=True)
                for issue in under15_issues:
                    note_key = f"note_{issue['row_idx']}"
                    current_note = st.session_state["dq_notes"].get(issue["row_idx"], "")
                    new_note = st.text_input(
                        f"سجل #{issue['row_idx']+1} — مسلسل {issue['serial']} — عمر: {issue['value']}",
                        value=current_note,
                        placeholder=t("dq_notes_placeholder"),
                        key=note_key
                    )
                    st.session_state["dq_notes"][issue["row_idx"]] = new_note

                if st.button(t("dq_save_notes"), key="save_dq_notes"):
                    st.success(t("dq_notes_saved"))

            # ── Export Quality Report ──
            st.markdown("---")
            if st.button(t("dq_export"), key="export_dq"):
                buf_dq = io.BytesIO()
                with pd.ExcelWriter(buf_dq, engine='openpyxl') as writer:
                    # Sheet 1: Errors
                    if errors:
                        err_rows = []
                        for issue in errors:
                            err_rows.append({
                                t("dq_col_row"):      issue["row_idx"] + 1,
                                "مسلسل":              issue["serial"],
                                t("dq_col_field"):    issue["field"],
                                t("dq_col_value"):    issue["value"],
                                t("dq_col_rule"):     t(issue["rule_key"]),
                                t("dq_col_severity"): "خطأ",
                                t("dq_col_notes"):    st.session_state["dq_notes"].get(issue["row_idx"], ""),
                            })
                        pd.DataFrame(err_rows).to_excel(writer, index=False, sheet_name="🔴 الأخطاء")

                    # Sheet 2: Warnings
                    if warnings:
                        warn_rows = []
                        for issue in warnings:
                            warn_rows.append({
                                t("dq_col_row"):      issue["row_idx"] + 1,
                                "مسلسل":              issue["serial"],
                                t("dq_col_field"):    issue["field"],
                                t("dq_col_value"):    issue["value"],
                                t("dq_col_rule"):     t(issue["rule_key"]),
                                t("dq_col_severity"): "تحذير",
                                t("dq_col_notes"):    st.session_state["dq_notes"].get(issue["row_idx"], ""),
                            })
                        pd.DataFrame(warn_rows).to_excel(writer, index=False, sheet_name="🟡 التحذيرات")

                    # Sheet 3: Summary
                    summary = pd.DataFrame([
                        {"البيان": "إجمالي السجلات",    "العدد": total},
                        {"البيان": "سجلات صحيحة 🟢",   "العدد": ok_count},
                        {"البيان": "أخطاء 🔴",          "العدد": len(errors)},
                        {"البيان": "تحذيرات 🟡",        "العدد": len(warnings)},
                        {"البيان": "تاريخ الفحص",       "العدد": datetime.now().strftime("%Y-%m-%d %H:%M")},
                    ])
                    summary.to_excel(writer, index=False, sheet_name="📊 الملخص")

                st.download_button(
                    "⬇️ تحميل تقرير الجودة",
                    data=buf_dq.getvalue(),
                    file_name=f"DataQuality_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_dq_report"
                )

# ──────────────────────────────────────
#  TAB 7 — INDICATORS DASHBOARD
# ──────────────────────────────────────
with tab_indicators:
    # ── Section 1: Indicators Dashboard ──
    st.markdown("""
    <div style="margin-bottom:1rem;">
      <div style="font-size:1.1rem;font-weight:700;color:#e0e0f0;">📈 Indicators Dashboard</div>
      <div style="font-size:0.78rem;color:#555;">الإنجاز الفعلي للمؤشرات البرامجية</div>
    </div>
    """, unsafe_allow_html=True)

    refusal_rate    = round((refused_test / total * 100), 1) if total > 0 else 0
    positivity_rate = round((positive     / total * 100), 1) if total > 0 else 0
    followup_rate   = round((has_followup / total * 100), 1) if total > 0 else 0
    referral_rate   = round((referrals    / positive * 100), 1) if positive > 0 else 0

    _ind_cols = st.columns(5)
    for _i, (_name, _val, _rate, _clr) in enumerate([
        ("HIV Tests",     total,            None,            "#7c6aff"),
        ("Positive",      positive,         positivity_rate, "#ff6b6b"),
        ("Referrals",     referrals,        referral_rate,   "#6bff8e"),
        ("Follow-up Visits", total_fu_visits, followup_rate, "#6bb3ff"),
        ("Refused",       refused_test,     refusal_rate,    "#ffb74d"),
    ]):
        with _ind_cols[_i]:
            st.markdown(f"""
            <div class="metric-card" style="border-color:{_clr}44;">
              <div class="label">{_name}</div>
              <div class="value" style="color:{_clr}">{_val:,}</div>
              {f'<div style="font-size:0.78rem;color:{_clr};font-weight:600;">{_rate}%</div>' if _rate is not None else ""}
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")

    # ── Section 2: Targets Management (Excel-style) ──
    st.markdown("""
    <div style="margin-bottom:0.8rem;">
      <div style="font-size:1rem;font-weight:700;color:#e0e0f0;">🎯 Targets Management</div>
      <div style="font-size:0.75rem;color:#555;">أدخل المستهدفات من Global Fund — Annual + Q1/Q2/Q3/Q4</div>
    </div>
    """, unsafe_allow_html=True)

    if "kpi_targets" not in st.session_state:
        st.session_state["kpi_targets"] = {
            k: {"annual": 0, "Q1": 0, "Q2": 0, "Q3": 0, "Q4": 0}
            for k in ["hiv_tests","positive","referrals","followups","refusal"]
        }

    _kpi_labels_t6 = {
        "hiv_tests":  "HIV Tests",
        "positive":   "Positive Cases",
        "referrals":  "Referrals",
        "followups":  "Follow-ups",
        "refusal":    "Refused",
    }
    _actuals_t6 = {
        "hiv_tests":  total,
        "positive":   positive,
        "referrals":  referrals,
        "followups":  total_fu_visits,
        "refusal":    refused_test,
    }

    # Header row
    _h0, _h1, _h2, _h3, _h4, _h5, _h6 = st.columns([1.8, 1, 1, 1, 1, 1, 1])
    for _col, _label in zip([_h0,_h1,_h2,_h3,_h4,_h5,_h6],
                             ["Indicator","Annual","Q1","Q2","Q3","Q4","Actual"]):
        _col.markdown(f"<div style='font-size:0.75rem;font-weight:700;color:#7c6aff;padding:0.3rem 0;border-bottom:1px solid #2a2a4e;'>{_label}</div>", unsafe_allow_html=True)

    # Data rows
    for _kk, _kname in _kpi_labels_t6.items():
        _r0,_r1,_r2,_r3,_r4,_r5,_r6 = st.columns([1.8,1,1,1,1,1,1])
        with _r0:
            st.markdown(f"<div style='font-size:0.82rem;color:#e0e0f0;padding:0.2rem 0;'>{_kname}</div>", unsafe_allow_html=True)
        for _q, _rc in zip(["annual","Q1","Q2","Q3","Q4"], [_r1,_r2,_r3,_r4,_r5]):
            with _rc:
                st.session_state["kpi_targets"][_kk][_q] = st.number_input(
                    "", min_value=0,
                    value=st.session_state["kpi_targets"][_kk][_q],
                    key=f"t6_{_kk}_{_q}", label_visibility="collapsed"
                )
        _actual_v = _actuals_t6[_kk]
        _ann_v    = st.session_state["kpi_targets"][_kk]["annual"]
        _pct_v    = round(_actual_v/_ann_v*100,1) if _ann_v > 0 else 0
        _c6 = "#6bff8e" if _pct_v>=90 else "#ffb74d" if _pct_v>=70 else "#ff6b6b"
        with _r6:
            st.markdown(f"""
            <div style='font-size:0.82rem;font-weight:700;color:{_c6};padding:0.4rem 0;'>
              {_actual_v:,}<br>
              <span style='font-size:0.7rem;'>{_pct_v}%</span>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")

    # ── Section 3: Achievement Chart ──
    _sel_q_t6 = st.selectbox(
        "عرض مقارنة:" if st.session_state["lang"]=="ar" else "Show comparison:",
        ["Q1","Q2","Q3","Q4","Annual"], key="sel_q_t6"
    )
    _q_key = "annual" if _sel_q_t6 == "Annual" else _sel_q_t6
    _chart_names = list(_kpi_labels_t6.values())
    _chart_act   = [_actuals_t6[k] for k in _kpi_labels_t6]
    _chart_tgt   = [st.session_state["kpi_targets"][k][_q_key] for k in _kpi_labels_t6]

    _fig_ind = go.Figure()
    _fig_ind.add_trace(go.Bar(name="Actual", x=_chart_names, y=_chart_act, marker_color="#7c6aff"))
    _fig_ind.add_trace(go.Bar(name=f"Target ({_sel_q_t6})", x=_chart_names, y=_chart_tgt, marker_color="#2a2a4e"))
    _fig_ind.update_layout(
        barmode="group", template="plotly_dark",
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(13,13,26,1)",
        title=f"Actual vs {_sel_q_t6} Target",
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
        height=350
    )
    st.plotly_chart(_fig_ind, use_container_width=True)

    # ── Donor Report ──
    st.markdown("---")
    st.markdown("### 📄 Donor Report")
    if st.button("📋 Generate Monthly Report", key="gen_report", use_container_width=False):
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile

        def _make_bar(labels, values, title, color="#7c6aff"):
            fig, ax = plt.subplots(figsize=(6,3.5), facecolor="white")
            ax.bar(labels, values, color=color, edgecolor="white")
            ax.set_title(title, fontsize=11, fontweight="bold")
            ax.set_facecolor("#f8f8f8")
            ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
            for b, v in zip(ax.patches, values):
                ax.text(b.get_x()+b.get_width()/2, b.get_height()+max(values)*0.01,
                        str(v), ha="center", fontsize=9, fontweight="bold")
            plt.tight_layout()
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            plt.savefig(tmp.name, dpi=150, bbox_inches="tight"); plt.close()
            return tmp.name

        def _build_doc(lang_code):
            doc = Document()
            for sec in doc.sections:
                sec.top_margin=Inches(1); sec.bottom_margin=Inches(1)
                sec.left_margin=Inches(1.2); sec.right_margin=Inches(1.2)
            is_ar = lang_code == "ar"
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("DataBridge — M&E Hub for Health Programs")
            r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(0x7c,0x6a,0xff)
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run("التقرير الشهري — Be Frienders 2026" if is_ar else "Monthly Report — Be Frienders 2026").font.size=Pt(13)
            p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.add_run(datetime.now().strftime("%B %Y")).font.size=Pt(11)
            doc.add_paragraph("─"*60)
            h = doc.add_heading("الملخص التنفيذي" if is_ar else "Executive Summary", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x7c,0x6a,0xff)
            doc.add_paragraph(
                f"إجمالي المستفيدين: {total:,} | إيجابي: {positive} ({positivity_rate}%) | إحالات: {referrals} | متابعة: {followup_rate}% | رفض: {refusal_rate}%"
                if is_ar else
                f"Beneficiaries: {total:,} | Positive: {positive} ({positivity_rate}%) | Referrals: {referrals} | Follow-up: {followup_rate}% | Refusal: {refusal_rate}%"
            )
            h2 = doc.add_heading("مؤشرات الأداء" if is_ar else "Key Performance Indicators", level=1)
            h2.runs[0].font.color.rgb = RGBColor(0x7c,0x6a,0xff)
            tbl = doc.add_table(rows=1+len(_kpi_labels_t6), cols=4)
            tbl.style = "Table Grid"
            for ci, ch in enumerate(["Indicator","Actual","Annual Target","Achievement"]):
                tbl.rows[0].cells[ci].text = ch
                tbl.rows[0].cells[ci].paragraphs[0].runs[0].bold = True
            for ri, (kk, kname) in enumerate(_kpi_labels_t6.items(), 1):
                act = _actuals_t6[kk]; ann = st.session_state["kpi_targets"][kk]["annual"]
                pct = f"{round(act/ann*100,1)}%" if ann>0 else "N/A"
                for ci, val in enumerate([kname, str(act), str(ann), pct]):
                    tbl.rows[ri].cells[ci].text = val
            doc.add_paragraph()
            h3 = doc.add_heading("الرسوم البيانية" if is_ar else "Charts", level=1)
            h3.runs[0].font.color.rgb = RGBColor(0x7c,0x6a,0xff)
            _ch = _make_bar(list(_kpi_labels_t6.values()), list(_actuals_t6.values()), "Actual Values")
            doc.add_picture(_ch, width=Inches(5.5))
            if qs:
                h4 = doc.add_heading("جودة البيانات" if is_ar else "Data Quality", level=1)
                h4.runs[0].font.color.rgb = RGBColor(0x7c,0x6a,0xff)
                doc.add_paragraph(f"Score: {qs['score']}% — {qs['status']} | Critical: {qs['critical']} | Major: {qs['major']}")
            doc.add_paragraph("─"*60)
            fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.add_run(f"Auto-generated by DataBridge — {datetime.now().strftime('%Y-%m-%d')}").font.size=Pt(9)
            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            return buf.getvalue()

        with st.spinner("⚙️ جاري إنشاء التقارير..."):
            _r_ar = _build_doc("ar"); _r_en = _build_doc("en")
        _rc1, _rc2 = st.columns(2)
        with _rc1:
            st.download_button("⬇️ التقرير العربي (.docx)", _r_ar,
                f"BeF_Report_AR_{datetime.now().strftime('%Y%m')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_r_ar")
        with _rc2:
            st.download_button("⬇️ English Report (.docx)", _r_en,
                f"BeF_Report_EN_{datetime.now().strftime('%Y%m')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_r_en")
        st.success("✅ تم إنشاء التقريرين!")



# ──────────────────────────────────────
#  AI SETTINGS MANAGER — Gemini model is configurable, not hard-coded
# ──────────────────────────────────────
GEMINI_STABLE_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
]
GEMINI_DEFAULT_MODEL = "gemini-2.5-flash"


def _get_ai_config() -> dict:
    cfg = load_config()
    cfg.setdefault("ai_provider", "gemini")
    cfg.setdefault("gemini_model", GEMINI_DEFAULT_MODEL)
    cfg.setdefault("gemini_api_key", "")
    return cfg


def _get_gemini_fallback_models(saved_model: str) -> list:
    ordered = [saved_model, "gemini-2.5-flash", "gemini-2.5-flash-lite"]
    out = []
    for m in ordered:
        m = str(m or "").strip()
        if m and m not in out:
            out.append(m)
    return out


def _call_gemini_model(api_key: str, model_name: str, prompt_text: str, max_tokens: int = 2048) -> str:
    import urllib.request
    import json as _jj

    payload = _jj.dumps({
        "contents": [{"parts": [{"text": prompt_text}]}],
        "generationConfig": {"temperature": 0.4, "maxOutputTokens": max_tokens}
    }).encode("utf-8")
    req = urllib.request.Request(
        f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}",
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=45) as resp:
        res = _jj.loads(resp.read().decode("utf-8"))
        return res["candidates"][0]["content"]["parts"][0]["text"]


def _call_gemini_with_fallbacks(api_key: str, saved_model: str, prompt_text: str) -> tuple:
    attempts = []
    for model_name in _get_gemini_fallback_models(saved_model):
        try:
            reply = _call_gemini_model(api_key, model_name, prompt_text)
            return reply, model_name, attempts
        except Exception as exc:
            attempts.append({"model": model_name, "error": str(exc)})
    details = "\n".join([f"- {a['model']}: {a['error']}" for a in attempts])
    raise RuntimeError("All configured Gemini models failed. Try another model in Settings.\n" + details)

# ──────────────────────────────────────
#  TAB 8 — AI ASSISTANT (Narrative)
# ──────────────────────────────────────
with tab_ai:
    _ai_cfg = _get_ai_config()
    _api_key = _ai_cfg.get("gemini_api_key", "")
    _gemini_model = _ai_cfg.get("gemini_model", GEMINI_DEFAULT_MODEL)

    st.markdown("""
    <div style="margin-bottom:1rem;">
      <div style="font-size:1.1rem;font-weight:700;color:#e0e0f0;">🤖 AI Data Assistant</div>
      <div style="font-size:0.78rem;color:#555;">مساعد ذكي متخصص في M&E — يكتب تقارير ويحلل البيانات</div>
    </div>
    """, unsafe_allow_html=True)

    if not _api_key:
        st.warning("⚠️ لم يتم إعداد Gemini API Key. اذهب إلى تاب ⚙️ Settings." if st.session_state["lang"]=="ar"
                   else "⚠️ Gemini API Key not configured. Go to ⚙️ Settings tab.")
    else:
        # ── Preset Questions ──
        st.markdown("#### 💡 أسئلة جاهزة" if st.session_state["lang"]=="ar" else "#### 💡 Quick Prompts")

        _presets_ar = [
            "📉 لماذا انخفضت نسبة المتابعة؟",
            "⚠️ ما المؤشرات المتأخرة عن المستهدف؟",
            "📝 اكتب تقريراً شهرياً للممول.",
            "🔧 اقترح إجراءات تصحيحية لتحسين الأداء.",
            "📊 حلّل توزيع الفئات العمرية وأبرز الملاحظات.",
            "🔍 ما أبرز مشاكل جودة البيانات وكيف تُحلّ؟",
        ]
        _presets_en = [
            "📉 Why has the follow-up rate declined?",
            "⚠️ Which indicators are behind target?",
            "📝 Write a monthly donor report.",
            "🔧 Suggest corrective actions to improve performance.",
            "📊 Analyze the age distribution and highlight key findings.",
            "🔍 What are the main data quality issues and how to fix them?",
        ]
        _presets = _presets_ar if st.session_state["lang"]=="ar" else _presets_en

        _pc = st.columns(3)
        for _pi, _pq in enumerate(_presets):
            with _pc[_pi % 3]:
                if st.button(_pq, key=f"preset_{_pi}", use_container_width=True):
                    st.session_state.setdefault("chat_history", [])
                    st.session_state["chat_history"].append({"role":"user","content":_pq})
                    st.session_state["ai_pending"] = _pq
                    st.rerun()

        st.markdown("---")

        # ── Chat History ──
        if "chat_history" not in st.session_state:
            st.session_state["chat_history"] = []

        for _msg in st.session_state["chat_history"]:
            _bg  = "#0d0d1a" if _msg["role"]=="assistant" else "#13132a"
            _brd = "#7c6aff44" if _msg["role"]=="assistant" else "#2a2a4e"
            _ico = "🤖" if _msg["role"]=="assistant" else "🧑"
            st.markdown(f"""
            <div style="background:{_bg};border:1px solid {_brd};border-radius:10px;
                        padding:0.9rem 1.1rem;margin:0.4rem 0;line-height:1.7;">
              <span style="font-size:0.75rem;color:#555;margin-left:0.4rem;">{_ico}</span>
              <span style="font-size:0.88rem;color:#e0e0f0;">{_msg["content"]}</span>
            </div>
            """, unsafe_allow_html=True)

        # ── Context Builder ──
        def _ctx():
            _lang = st.session_state["lang"]
            _kk_map = {"hiv_tests":"HIV Tests","positive":"Positive","referrals":"Referrals","followups":"Follow-ups","refusal":"Refused"}
            _kpi_lines = "\n".join([
                f"- {_kk_map[k]}: Actual={[total,positive,referrals,has_followup,refused_test][i]}, Annual Target={st.session_state.get('kpi_targets',{}).get(k,{}).get('annual',0)}"
                for i,k in enumerate(_kk_map)
            ])
            _qs_txt = f"Quality Score: {qs['score']}% ({qs['status']}), Critical={qs['critical']}, Major={qs['major']}" if qs else "Not run"
            return f"""You are an expert M&E analyst and report writer for HIV prevention programs (Be Frienders 2026, Egypt, Global Fund).
Respond in {"Arabic" if _lang=="ar" else "English"}.
When asked for a report or narrative: write in professional paragraphs, use bullet points for findings, and include actionable recommendations.
Be analytical, precise, and cite actual numbers.

PROGRAM DATA:
- Total Beneficiaries: {total:,}
- Positive Cases: {positive} ({positivity_rate}%)
- Referrals to Treatment: {referrals} ({referral_rate}% of positives)
- Follow-up Rate: {followup_rate}% ({has_followup:,} beneficiaries)
- Refusal Rate: {refusal_rate}%

KPI vs TARGETS:
{_kpi_lines}

DATA QUALITY: {_qs_txt}"""

        # ── Process pending AI request ──
        if st.session_state.get("ai_pending"):
            _q = st.session_state.pop("ai_pending")
            with st.spinner("🤖 جاري التحليل..." if st.session_state["lang"]=="ar" else "🤖 Analyzing..."):
                try:
                    _prompt_text = _ctx() + f"\n\nUSER REQUEST:\n{_q}"
                    _reply, _used_model, _attempts = _call_gemini_with_fallbacks(_api_key, _gemini_model, _prompt_text)
                    if _used_model != _gemini_model:
                        _reply = f"⚠️ الموديل المحفوظ ({_gemini_model}) غير متاح حالياً، تم استخدام {_used_model} تلقائياً.\n\n" + _reply
                        _ai_cfg["gemini_model"] = _used_model
                        save_config(_ai_cfg)
                except Exception as _e:
                    _reply = f"❌ {_e}"
            st.session_state["chat_history"].append({"role":"assistant","content":_reply})
            st.rerun()

        # ── Chat Input ──
        _user_q = st.chat_input(
            "اكتب سؤالك أو طلبك..." if st.session_state["lang"]=="ar"
            else "Type your question or request..."
        )
        if _user_q:
            st.session_state["chat_history"].append({"role":"user","content":_user_q})
            st.session_state["ai_pending"] = _user_q
            st.rerun()

        if st.session_state.get("chat_history"):
            if st.button("🗑️ مسح المحادثة" if st.session_state["lang"]=="ar" else "🗑️ Clear Chat", key="clear_chat"):
                st.session_state["chat_history"] = []
                st.rerun()


# ──────────────────────────────────────
#  TAB 9 — SETTINGS
# ──────────────────────────────────────
with tab_settings:
    st.markdown("""
    <div style="margin-bottom:1.5rem;">
      <div style="font-size:1.1rem;font-weight:700;color:#e0e0f0;">⚙️ Settings</div>
      <div style="font-size:0.78rem;color:#555;">إعدادات البرنامج — تُحفظ تلقائياً في AppData / DataBridge</div>
    </div>
    """, unsafe_allow_html=True)

    cfg2 = load_config()
    APPDATA_DIR2 = get_appdata_dir()

    st.markdown("### 🔑 AI Settings Manager")
    st.markdown(f"""
    <div class="info-box">
    احصل على مفتاحك المجاني من:
    <a href="https://aistudio.google.com/app/apikey" target="_blank" style="color:#7c6aff;">
    Google AI Studio
    </a>
    <br>الإعدادات تُحفظ في: <code style="color:#7c6aff;">{APPDATA_DIR2}</code>
    </div>
    """, unsafe_allow_html=True)

    cfg2.setdefault("ai_provider", "gemini")
    cfg2.setdefault("gemini_model", GEMINI_DEFAULT_MODEL)
    cfg2.setdefault("gemini_api_key", "")

    current_key = cfg2.get("gemini_api_key", "")
    masked = f"{'*' * (len(current_key)-4)}{current_key[-4:]}" if len(current_key) > 4 else ""
    if masked:
        st.markdown(f'<div style="font-size:0.8rem;color:#6bff8e;margin-bottom:0.5rem;">✅ مفتاح محفوظ: <code>{masked}</code></div>', unsafe_allow_html=True)

    st.selectbox(
        "AI Provider" if st.session_state["lang"] == "en" else "مزود الذكاء الاصطناعي",
        ["gemini"], index=0, key="settings_ai_provider", disabled=True,
        help="حالياً Gemini فقط، ويمكن إضافة OpenAI أو Local LLM لاحقاً."
    )

    saved_model = str(cfg2.get("gemini_model", GEMINI_DEFAULT_MODEL)).strip() or GEMINI_DEFAULT_MODEL
    model_choices = ["gemini-2.5-flash", "gemini-2.5-flash-lite", "Custom model name"]
    default_idx = model_choices.index(saved_model) if saved_model in model_choices else 2
    selected_model_choice = st.selectbox(
        "Gemini Model" if st.session_state["lang"] == "en" else "موديل Gemini",
        model_choices, index=default_idx, key="settings_gemini_model_choice"
    )
    custom_model = ""
    if selected_model_choice == "Custom model name":
        custom_model = st.text_input(
            "Custom Gemini model name" if st.session_state["lang"] == "en" else "اسم موديل Gemini مخصص",
            value=saved_model if saved_model not in model_choices else "",
            placeholder="مثال: gemini-2.5-flash",
            key="settings_gemini_custom_model"
        )
    chosen_model = custom_model.strip() if selected_model_choice == "Custom model name" else selected_model_choice

    new_key = st.text_input(
        "أدخل Gemini API Key" if st.session_state["lang"] == "ar" else "Enter Gemini API Key",
        type="password", placeholder="AIza...", key="settings_api_key"
    )

    c_save_ai, c_test_ai = st.columns(2)
    with c_save_ai:
        if st.button("💾 حفظ إعدادات AI" if st.session_state["lang"] == "ar" else "💾 Save AI Settings", key="save_ai_settings", use_container_width=True):
            if not chosen_model:
                st.error("❌ اختر أو اكتب اسم موديل صحيح" if st.session_state["lang"] == "ar" else "❌ Please choose or enter a valid model name")
            else:
                cfg2["ai_provider"] = "gemini"
                cfg2["gemini_model"] = chosen_model
                if new_key.strip():
                    cfg2["gemini_api_key"] = new_key.strip()
                save_config(cfg2)
                st.success("✅ تم حفظ إعدادات AI في AppData" if st.session_state["lang"] == "ar" else "✅ AI settings saved to AppData")
                st.rerun()
    with c_test_ai:
        if st.button("🧪 Test AI Connection" if st.session_state["lang"] == "en" else "🧪 اختبار اتصال AI", key="test_ai_connection", use_container_width=True):
            test_key = new_key.strip() or cfg2.get("gemini_api_key", "")
            if not test_key:
                st.error("❌ أدخل أو احفظ Gemini API Key أولاً" if st.session_state["lang"] == "ar" else "❌ Enter or save Gemini API Key first")
            elif not chosen_model:
                st.error("❌ اختر أو اكتب اسم موديل صحيح" if st.session_state["lang"] == "ar" else "❌ Choose or enter a valid model name")
            else:
                with st.spinner("جاري اختبار الاتصال..." if st.session_state["lang"] == "ar" else "Testing connection..."):
                    try:
                        _test_reply, _used_model, _attempts = _call_gemini_with_fallbacks(test_key, chosen_model, "Reply with OK only.")
                        cfg2["ai_provider"] = "gemini"
                        cfg2["gemini_model"] = _used_model
                        if new_key.strip():
                            cfg2["gemini_api_key"] = new_key.strip()
                        save_config(cfg2)
                        if _used_model == chosen_model:
                            st.success(f"✅ AI Connection OK — Model: {_used_model}")
                        else:
                            st.warning(f"⚠️ الموديل المختار غير متاح، لكن الاتصال نجح باستخدام fallback: {_used_model}")
                    except Exception as _e:
                        st.error("❌ Model not available. جرّب اختيار موديل آخر من الإعدادات.")
                        st.caption(str(_e))

    st.markdown("---")
    st.markdown("### ℹ️ معلومات البرنامج")
    st.markdown(f"""
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;">
      <div class="metric-card">
        <div class="label">Version</div>
        <div style="font-size:1rem;font-weight:600;color:#7c6aff;">{APP_VERSION}</div>
      </div>
      <div class="metric-card">
        <div class="label">Config Path</div>
        <div style="font-size:0.7rem;color:#888;word-break:break-all;">{APPDATA_DIR2}</div>
      </div>
      <div class="metric-card">
        <div class="label">Users File</div>
        <div style="font-size:0.75rem;color:#888;">{USERS_FILE}</div>
      </div>
      <div class="metric-card">
        <div class="label">Max Upload</div>
        <div style="font-size:1rem;font-weight:600;color:#7c6aff;">{MAX_UPLOAD_SIZE_MB} MB</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("🗑️ حذف API Key" if st.session_state["lang"] == "ar" else "🗑️ Delete API Key",
                 key="delete_api_key"):
        cfg2["gemini_api_key"] = ""
        save_config(cfg2)
        st.success("✅ تم حذف المفتاح" if st.session_state["lang"] == "ar" else "✅ Key deleted")
        st.rerun()

# ──────────────────────────────────────
#  TAB 10 — SOURCE CONVERTER
# ──────────────────────────────────────
with tab_converter:
    st.markdown("""
    <div style="margin-bottom:1.5rem;">
      <div style="font-size:1.1rem;font-weight:700;color:#e0e0f0;">🔄 Source Converter</div>
      <div style="font-size:0.78rem;color:#555;">تحويل الشيت الأساسي (204 عمود) إلى شيت البرنامج (27 عمود) تلقائياً</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box">
    ارفع الشيت الأساسي (المكوّن من 204 عمود) وسيتحول تلقائياً إلى هيكل شيت البرنامج جاهزاً للتحليل.
    </div>
    """, unsafe_allow_html=True)

    _src_file = st.file_uploader(
        "ارفع الشيت الأساسي (Excel)",
        type=["xlsx"], key="src_converter_upload"
    )

    if _src_file:
        try:
            validate_uploaded_file(_src_file)
            _src_bytes = _src_file.read()

            # Preview source
            _df_prev = pd.read_excel(io.BytesIO(_src_bytes), header=None, nrows=2)
            st.markdown(f'<div class="info-box">📄 الملف: <b>{_src_file.name}</b> — {len(_df_prev.columns)} عمود</div>', unsafe_allow_html=True)

            if st.button("🔄 تحويل الآن", key="run_convert", use_container_width=False):
                with st.spinner("⚙️ جاري التحويل..."):

                    import warnings as _w
                    _w.filterwarnings("ignore")

                    _df_src = pd.read_excel(io.BytesIO(_src_bytes), header=None, skiprows=4)
                    _df_src.columns = range(len(_df_src.columns))

                    # ── وقّف عند آخر صف فيه كود مجمع أو تاريخ زيارة ──
                    _lc = _df_src[15].last_valid_index()
                    _ld = _df_src[2].last_valid_index()
                    _lr = max(x for x in [_lc, _ld] if x is not None)
                    _df_src = _df_src.iloc[:_lr + 1].copy()

                    def _col(i):
                        return _df_src[i-1] if (i-1) in _df_src.columns else pd.Series([None]*len(_df_src))

                    def _yes_qty(s_yn, s_qty):
                        res = []
                        for yn, qty in zip(s_yn, s_qty):
                            if str(yn).strip() == 'نعم':
                                try:    res.append(int(float(qty)) if pd.notna(qty) else None)
                                except: res.append(None)
                            else:
                                res.append(None)
                        return res

                    def _yn_simple(s):
                        return s.apply(lambda x: 'نعم' if str(x).strip() == 'نعم' else ('لا' if pd.notna(x) and str(x).strip() not in ['nan',''] else None))

                    _out = pd.DataFrame()
                    _out['مسلسل']                                          = _col(1)
                    _out['الكود المجمع']                                   = _col(16)
                    _out['اسم الجمعية']                                    = _col(2)
                    _out['تاريخ الزيارة']                                  = format_idus_source_date(_col(3), swap_excel_dates=False)
                    _out['اسم مندوب الوصول ثلاثي\n(الباحث الميداني) ']   = _col(4)
                    _out['محافظة الوصول للمستفيد']                        = _col(5)
                    _out['منطقة الوصول للمستفيد ']                        = _col(6)
                    _out['مكان المقابلة']                                  = _col(17)
                    _out['السن']                                           = _col(18)
                    _out['النوع']                                          = _col(19)
                    _out['واقيات']                                         = _yes_qty(_col(27), _col(28))
                    _out['مزلقات']                                         = _yes_qty(_col(31), _col(32))
                    _out[' زهري']                                          = _yn_simple(_col(34))
                    _out['دعم نفسي']                                       = _yn_simple(_col(36))
                    _out['سرنجات']                                         = _yes_qty(_col(39), _col(40))
                    _out['ميثادون']                                        = _yn_simple(_col(42))
                    _out['نتيجة التحليل ']                                 = _col(48)
                    _out['نتيجة التحليل التاكيدي']                        = _col(50)
                    _out['الاحالة علي العلاج']                            = _col(51)
                    _out['زيارة متابعة 1\n(تاريخ الزيارة)']              = format_idus_source_date(_col(57))
                    _out['واقيات متابعه']                                  = _yes_qty(_col(62), _col(63))
                    _out['مزلقات متابعه']                                  = _yes_qty(_col(66), _col(67))
                    _out['زهري متابعه']                                    = _yn_simple(_col(69))
                    _out['الدعم النفسي متابعه']                           = _yn_simple(_col(71))
                    _out['سرنجات متابعه']                                  = _yes_qty(_col(74), _col(75))
                    _out['نتيجة التحليل متابعه']                          = _col(78)
                    _out['ميثادون متابعه']                                 = _yn_simple(_col(80))

                    _out = _out.dropna(how='all').reset_index(drop=True)
                    _out.columns = [str(c).strip() for c in _out.columns]
                    _out, _conv_report = clean_dataframe(_out)
                    _conv_standard = generate_standard_df(_out)

                    _buf_out = io.BytesIO()
                    with pd.ExcelWriter(_buf_out, engine='openpyxl') as _wr:
                        _out.to_excel(_wr, index=False, sheet_name='IDUs Database')
                    _buf_out.seek(0)
                    _result_bytes = _buf_out.getvalue()

                st.success(f"✅ تم التحويل والتنظيف! {len(_out):,} سجل — {len(_out.columns)} عمود")

                _conv_map_df = mapping_report_to_dataframe(_conv_report)
                if not _conv_map_df.empty:
                    st.markdown("**Auto Mapper للملف المحوّل — راجع الربط قبل الاعتماد**")
                    st.dataframe(_conv_map_df, use_container_width=True, height=220)

                if st.button("✅ استخدام الملف المحوّل في التحليل وفتح شاشة الاعتماد", key="use_converted_for_analysis"):
                    st.session_state["hdf"] = _out
                    st.session_state["data_clean_report"] = _conv_report
                    st.session_state["standard_df"] = _conv_standard
                    st.session_state["mapper_approved"] = False
                    st.session_state["dq_issues"] = None
                    st.session_state["hiv_file_name"] = f"[محوّل] {_src_file.name}"
                    st.rerun()

                # Preview
                st.dataframe(_out.head(10), use_container_width=True, height=280)

                st.download_button(
                    "⬇️ تحميل الشيت المحوّل",
                    data=_result_bytes,
                    file_name=f"IDUs_Converted_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_converted"
                )

        except Exception as _ce:
            st.error(f"❌ خطأ في التحويل: {_ce}")


# ──────────────────────────────────────
#  TAB 2 — DATA REPAIR CENTER (shown as second tab)
# ──────────────────────────────────────
with tab_repair:
    repaired_df = render_data_repair_center(st.session_state["hdf"])
    if repaired_df is not None:
        st.session_state["hdf"] = repaired_df
        st.session_state["standard_df"] = generate_standard_df(repaired_df)
        st.session_state["dq_issues"] = None
        st.session_state["mapper_approved"] = True
        st.info("تم تحديث البيانات داخل الجلسة. سيتم إعادة تحميل التحليل الآن.")
        st.rerun()
